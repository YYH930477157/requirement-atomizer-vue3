from __future__ import annotations

import io
import json
import tempfile
import unittest
from contextlib import redirect_stderr, redirect_stdout
from pathlib import Path

from docx import Document

import desktop_tasks
from api_server import build_review_summary
from io_utils import read_jsonl
from result_package import load_result_package, resolve_analysis_root
from review_state import apply_expert_decision


class ResultPackageEndToEndTests(unittest.TestCase):
    def test_analysis_completion_and_review_survive_clean_package_reload(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            source = base / "meter.docx"
            result = base / "result"
            document = Document()
            document.add_heading("Meter requirements", level=1)
            document.add_paragraph(
                "The meter shall support configurable auxiliary outputs."
            )
            document.save(source)

            sink = io.StringIO()
            with redirect_stdout(sink), redirect_stderr(sink):
                self.assertEqual(desktop_tasks.main([
                    "result-package-start", "--out", str(result),
                    "--input", str(source), "--stages", "atomize,ai-extract",
                ]), 0)
                run_id = load_result_package(result)["active_attempt"]["run_id"]
                self.assertEqual(desktop_tasks.main([
                    "run", "--input", str(source), "--out", str(result),
                    "--skip-review",
                ]), 0)
                self.assertEqual(desktop_tasks.main([
                    "chain", "--out", str(result), "--stages", "ai-extract",
                    "--llm-route", "stub",
                ]), 0)
                self.assertEqual(desktop_tasks.main([
                    "result-package-complete", "--out", str(result),
                    "--run-id", run_id,
                    "--completed-stages", "atomize,ai-extract",
                ]), 0)

            package = load_result_package(result, verify=True)
            self.assertEqual(package["analysis_status"], "completed")
            root_names = {path.name for path in result.iterdir()}
            self.assertEqual(
                root_names,
                {".ratomizer", "result-package.json", "summary.md", "merged_spec.xlsx"},
            )
            analysis_root = resolve_analysis_root(result)
            state_root = result / ".ratomizer" / "state"
            self.assertTrue((state_root / "claim_generation.meta.json").is_file())
            self.assertTrue((state_root / "claim_effective_ledger.jsonl").is_file())
            self.assertFalse((analysis_root / "claim_generation.meta.json").exists())

            requirement = read_jsonl(analysis_root / "atomic_requirements.jsonl")[0]
            requirement_id = str(
                requirement.get("stable_req_id")
                or requirement.get("requirement_id")
                or requirement.get("req_id")
            )
            marker_before_review = (result / "result-package.json").read_bytes()
            apply_expert_decision(
                analysis_root,
                requirement_id,
                "accepted",
                actor="expert",
                reason="E2E review persistence",
            )

            self.assertTrue((state_root / "review_states.jsonl").is_file())
            self.assertFalse((analysis_root / "review_states.jsonl").exists())
            self.assertEqual(
                (result / "result-package.json").read_bytes(), marker_before_review
            )
            summary = build_review_summary(resolve_analysis_root(result))
            self.assertEqual(summary["status_counts"].get("accepted"), 1)
            self.assertEqual(load_result_package(result)["analysis_status"], "completed")


if __name__ == "__main__":
    unittest.main()
