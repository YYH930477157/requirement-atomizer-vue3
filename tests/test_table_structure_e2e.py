"""表格结构闭环端到端回归（F10）。

- 真实 docx 映射矩阵 → atomize → catalog → publish → fold/queue → reload →
  execute 全链真实驱动（仅 LLM chat 注入，其余零 mock 核心链路）；
- xlsx sheet 级非空格守恒计数器：区域拆分漏格/重叠 → parse_incomplete 硬失败。

P1-4 诚实化：夹具之一（Feature/Behavior/Note + 义务句格）实际分类为
other/mixed 而非 mapping_matrix——其无信号自由文本格（"see below"）按
P1-1/B5 如实计数 unsignaled 并置 needs_review，不再冒充 status=ok；
另补真正的 marker mapping_matrix 夹具（GET×SET 矩阵）。模拟覆盖验证器
逐项推导七维检查（行头主体/列头维度/情态/极性/数量/条件/成文义务），
不再只对 modal 片段却全置 true；模拟传输的付费调用数与终态 usage 窗口
必须逐次相等（崩溃 usage 窗口的进程级探针见
tests/test_claim_queue_execution.py::test_crash_after_base_publication_recovers_full_verifier_usage）。
"""
from __future__ import annotations

import copy
import json
import os
import re
import tempfile
import unittest
from pathlib import Path
from typing import Any
from unittest import mock

from docx import Document
from jsonschema import Draft202012Validator, ValidationError
from openpyxl import Workbook
from openpyxl.worksheet.table import Table as ExcelTable

import ai_extract
import claim_artifacts
import claim_catalog
import claim_ledger
import claim_queue_execution as execution
import claim_review_actions
import doc_annotation_export as dae
import llm_client
import omission_actions
from atomize import extract_docx
from claim_artifacts import atomic_write_jsonl, file_sha256
from parsers.xlsx_parser import extract_xlsx
from requirement_kb import KnowledgeRepository


def _write_matrix_docx(path: Path) -> None:
    document = Document()
    document.add_heading("5 Requirements", level=1)
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Behavior"
    table.cell(0, 2).text = "Note"
    table.cell(1, 0).text = "Encryption"
    # 多义务格（F8）：一格两句独立规范性句 → 按句切出两个 table_cell claim
    table.cell(1, 1).text = (
        "The meter shall authenticate all clients. "
        "The meter shall log authentication failures."
    )
    table.cell(1, 2).text = "see below"
    table.cell(2, 0).text = "Signing"
    # 代词主体（P1-4）：句内无具体主体，主体身份只能来自行头上下文——
    # 模拟验证器必须真的核对行头，而不是对 modal 片段全置 true
    table.cell(2, 1).text = "It shall sign responses."
    table.cell(2, 2).text = "free text"
    document.save(path)


def _write_marker_matrix_docx(path: Path) -> None:
    """真正的 mapping_matrix（P1-4）：marker 格 × 正向维度列头。"""
    document = Document()
    document.add_heading("5 Requirements", level=1)
    document.add_paragraph(
        "Table 1 - Optical communication interface capabilities"
    )
    table = document.add_table(rows=3, cols=3)
    for row_index, row in enumerate((
        ("Interface", "GET", "SET"),
        ("Data access", "X", "X"),
        ("Event push", "X", ""),
    )):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    document.save(path)


# ---------------------------------------------------------------------------
# P1-4 诚实模拟层：一切响应从真实请求 prompt 派生；七维检查逐项独立判定，
# 行头主体/列头维度缺失时如实 covered=false（不再对 modal 片段全置 true）。
# ---------------------------------------------------------------------------

def _norm_text(value: str) -> str:
    return " ".join(str(value or "").split()).casefold()


_MODAL_RE = re.compile(r"\b(?:shall|must)\b", re.IGNORECASE)
_NEGATION_RE = re.compile(r"\b(?:not|never|no|without)\b", re.IGNORECASE)
_CONDITION_RE = re.compile(r"\b(?:if|when|unless|whenever)\b[^.]*", re.IGNORECASE)
_NUMBER_RE = re.compile(r"\d+(?:\.\d+)?")


def _parse_claim_context(source_claim: str) -> dict[str, Any]:
    """cell claim 的 semantic_context → 表标题/身份对/维度/正文分句。"""
    segments = [
        segment.strip()
        for segment in str(source_claim or "").split(" | ")
        if segment.strip()
    ]
    identities: list[tuple[str, str]] = []
    dimension, body = "", ""
    if segments:
        last = segments[-1]
        if " = " in last:
            dimension, body = (part.strip() for part in last.split(" = ", 1))
        else:
            body = last
        for segment in segments[1:-1]:
            if "=" in segment:
                key, value = (part.strip() for part in segment.split("=", 1))
                identities.append((key, value))
    return {
        "title": segments[0] if segments else "",
        "identities": identities,
        "subject": identities[-1][1] if identities else "",
        "dimension": dimension,
        "body": body,
    }


def _honest_coverage_decision(
    source_claim: str,
    evidence_text: str,
) -> tuple[bool, list[bool]]:
    """逐项推导七维检查；covered ⟺ 全部通过（与账本一致性规则同构）。"""
    context = _parse_claim_context(source_claim)
    body = str(context["body"] or "")
    body_norm = _norm_text(body)
    union = _norm_text(evidence_text)
    is_marker = bool(body_norm) and len(body_norm) <= 3 and body_norm.isalnum()
    modal_in_evidence = bool(_MODAL_RE.search(union))
    checks = {
        # 行头主体必须真实落入证据（代词/裸 marker 主体不可白送）
        "subject": bool(context["subject"])
        and _norm_text(str(context["subject"])) in union,
        "modality": (
            modal_in_evidence
            if is_marker
            else bool(_MODAL_RE.search(body)) and modal_in_evidence
        ),
        "polarity": bool(_NEGATION_RE.search(body_norm))
        == bool(_NEGATION_RE.search(union)),
        "quantities_units": all(
            number in union for number in _NUMBER_RE.findall(body)
        ),
        "conditions_exceptions": all(
            _norm_text(fragment) in union
            for fragment in _CONDITION_RE.findall(body)
        ),
        # 列头维度（表标题/列上下文）必须真实落入证据
        "scope": bool(context["dimension"])
        and _norm_text(str(context["dimension"])) in union,
        "target_obligation_framing": modal_in_evidence,
    }
    covered = all(checks.values())
    return covered, [checks[name] for name in claim_ledger.SEMANTIC_COVERAGE_CHECKS]


def _prompt_section_items(user: str, marker: str) -> list[str]:
    """解析 P0-3 结构化 prompt 小节（「- 」条目，空行/非条目行终止）。"""
    if marker not in user:
        return []
    tail = user.split(marker, 1)[1]
    if "\n" not in tail:
        return []
    items: list[str] = []
    for line in tail.split("\n", 1)[1].splitlines():
        if line.startswith("- "):
            items.append(line[2:].strip())
        else:
            break
    return items


def _simulated_extraction(user: str) -> str:
    """严格定向抽取的模拟响应：只从 prompt 证据段派生，缺证据如实空数组。

    产出一律绑定行头主体 + 列头维度（模拟验证器会真实核对这些维度）；
    上下文段（定位上下文）本身永不充当 source_quote。
    """
    requirements: list[dict[str, Any]] = []
    context_items = _prompt_section_items(user, "定位上下文（")
    table_title = ""
    subject = ""
    dimension = ""
    for item in context_items:
        if item.startswith("表标题："):
            table_title = item.split("：", 1)[1].strip()
        elif item.startswith("列头："):
            dimension = item.split("：", 1)[1].strip()
        elif item.startswith("行头："):
            subject = item.split("：", 1)[1].rsplit("=", 1)[-1].strip()
        elif "=" in item:  # v2 prompt compatibility for focused helper probes
            subject = item.rsplit("=", 1)[-1].strip()
        elif not dimension:
            dimension = item
    scope_prefix = f"{table_title} — " if table_title else ""
    prefix = (
        f"{scope_prefix}{subject} — {dimension}: "
        if subject and dimension
        else scope_prefix
    )
    for line in _prompt_section_items(user, "可引用证据（"):
        obligations = [
            fragment.strip()
            for fragment in re.split(r"(?<=[.!?])\s+", line)
            if _MODAL_RE.search(fragment)
        ]
        for quote in obligations[:1]:
            requirements.append({
                "title": " ".join(quote.split()[:8]).rstrip("."),
                "description": f"{prefix}{quote}",
                "source_quote": quote,
                "module": "security",
                "sub_items": [],
                "acceptance_criteria": [],
            })
    for item in _prompt_section_items(user, "矩阵事实（"):
        match = re.search(r"矩阵事实：主体=(.*?)\s*｜\s*维度=(.*?)\s*｜\s*取值=(.*)$", item)
        if not match:
            continue
        raw_subject, fact_dimension, marker = (
            part.strip() for part in match.groups()
        )
        fact_subject = raw_subject.rsplit("=", 1)[-1].strip()
        requirements.append({
            "title": f"{fact_subject} {fact_dimension}",
            "description": (
                f"{scope_prefix}{fact_subject} — {fact_dimension}: "
                "the product shall support "
                f"this capability (matrix marker {marker})."
            ),
            "source_quote": marker,
            "module": "security",
            "sub_items": [],
            "acceptance_criteria": [],
        })
    return json.dumps(
        {"requirements": requirements, "supplements": []}, ensure_ascii=False
    )


def _simulated_paid_content(payload: dict[str, Any]) -> str:
    """按请求类型分发模拟响应；非预期 prompt 响亮报错（付费面收敛）。"""
    messages = payload.get("messages") or []
    user = next(
        (
            str(message.get("content") or "")
            for message in reversed(messages)
            if message.get("role") == "user"
        ),
        "",
    )
    if "claim-coverage-verifier-request/v2" in user:
        # 独立覆盖验证器：七维逐项从 claim 上下文与 target evidence 实文推导
        request = json.loads(user)
        evidence = request.get("target_evidence") or []
        decisions = []
        for group in request.get("groups") or []:
            group_ref, source_claim, target_refs = group
            union = " ".join(
                str(text)
                for ref in target_refs
                for text in (evidence[ref] or [])
            )
            covered, checks = _honest_coverage_decision(str(source_claim or ""), union)
            decisions.append([group_ref, covered, checks])
        return json.dumps({"decisions": decisions}, ensure_ascii=False)
    if "claim-negative-proposer-request/v1" in user:
        # 本文档全部为规范性内容，无 semantic-negative 候选可提
        return json.dumps({"proposals": []}, ensure_ascii=False)
    if "claim-negative-verifier-request/v1" in user:
        request = json.loads(user)
        decisions = [
            {
                "claim_id": str(row.get("claim_id") or ""),
                "non_normative": False,
                "checks": {
                    **{
                        name: True
                        for name in claim_ledger.SEMANTIC_NEGATIVE_CHECKS
                    },
                    "no_normative_obligation": False,
                },
                "reason": "definition",
                "evidence": [],
                "rationale": "obligation sentence stays normative",
            }
            for row in request.get("claims") or []
        ]
        return json.dumps({"decisions": decisions}, ensure_ascii=False)
    if "【严格定向需求抽取】" in user:
        return _simulated_extraction(user)
    raise RuntimeError(f"unexpected paid prompt in e2e: {user[:200]}")


def _make_fake_post_json(recorded_payloads: list[dict[str, Any]]):
    """复刻真实 _post_json 的付费契约：网络前 reserve、成功 commit、stats 计数。"""

    def fake_post_json(
        config,
        payload,
        *,
        _request_budget=None,
        _request_stats=None,
    ):
        reservation_id = (
            _request_budget.reserve(payload)
            if _request_budget is not None
            else None
        )
        if _request_stats is not None:
            _request_stats["call_count"] = int(
                _request_stats.get("call_count") or 0
            ) + 1
        recorded_payloads.append(copy.deepcopy(payload))
        try:
            content = _simulated_paid_content(payload)
        except Exception:
            if _request_budget is not None and reservation_id is not None:
                _request_budget.fail(reservation_id)
            raise
        usage = {
            "prompt_tokens": 64,
            "completion_tokens": 32,
            "total_tokens": 96,
        }
        if _request_budget is not None and reservation_id is not None:
            _request_budget.commit(reservation_id, usage)
        return {
            "choices": [{
                "message": {"role": "assistant", "content": content},
                "finish_reason": "stop",
            }],
            "usage": usage,
        }

    return fake_post_json


_ROUTE_ENV = {
    "RATOMIZER_LLM_BASE_URL": "https://example.invalid/v1",
    "RATOMIZER_LLM_MODEL": "deepseek-chat",
    "RATOMIZER_LLM_MAX_RETRIES": "0",
    "RATOMIZER_LLM_MAX_TOKENS": "128",
}


def _assert_paid_surface_closed(test_case, recorded_payloads: list[dict[str, Any]]) -> None:
    """付费面收敛：除定向抽取与 claim 验证器请求外无任何意外付费调用。"""
    for payload in recorded_payloads:
        users = [
            str(message.get("content") or "")
            for message in (payload.get("messages") or [])
            if message.get("role") == "user"
        ]
        test_case.assertTrue(
            any(
                "【严格定向需求抽取】" in content
                or "claim-coverage-verifier-request/v2" in content
                or "claim-negative-proposer-request/v1" in content
                or "claim-negative-verifier-request/v1" in content
                for content in users
            ),
            f"unexpected paid payload: {users[:1]}",
        )



class TableStructureEndToEndTests(unittest.TestCase):
    """catalog → queue → publish → reload → execute 闭环（真实 docx）。"""

    def test_catalog_queue_publish_reload_execute_loop(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx_path = root / "matrix.docx"
            _write_matrix_docx(docx_path)
            out = root / "out"
            out.mkdir()

            # ---- 1) atomize（真实 docx 解析） --------------------------------
            blocks, items, cells = extract_docx(docx_path)
            atomic_write_jsonl(out / "blocks.jsonl", blocks)
            atomic_write_jsonl(out / "table_items.jsonl", items)
            atomic_write_jsonl(out / "table_cell_items.jsonl", cells)

            # ---- 2) catalog + 守恒矩阵 ----------------------------------------
            catalog = claim_catalog.build_claim_catalog(
                blocks, items, table_cell_items=cells
            )
            meta = catalog["meta"]
            self.assertEqual(meta["accounting_status"], "complete")
            # P1-4：本夹具实为 other/mixed——"see below" 无信号自由文本格如实
            # 计数 unsignaled 并置 needs_review，不再冒充 status=ok 的映射矩阵
            self.assertEqual(meta["table_structure_status"], "needs_review")
            audit = meta["audit"]
            self.assertEqual(audit.get("unsignaled_data_cell_count") or 0, 1)
            for key in (
                "unconsumed_table_cell_count",
                "multi_consumed_table_cell_count",
                "dangling_table_cell_reference_count",
                "unconsumed_table_item_count",
                "multi_consumed_table_item_count",
                "normative_context_only_count",
                "duplicate_table_cell_id_count",
            ):
                self.assertEqual(audit.get(key) or 0, 0, f"audit counter {key}")
            cell_claims = [
                row for row in catalog["catalog"]
                if row["source_kind"] == "table_cell"
                and row["eligibility"] == "claim"
            ]
            excluded_cell_candidate = next(
                row for row in catalog["catalog"]
                if row["source_kind"] == "table_cell"
                and row["eligibility"] == "excluded"
                and row["text"] == "see below"
            )
            self.assertEqual(
                excluded_cell_candidate["exclusion"]["reason"],
                "unsignaled_table_cell",
            )
            # 多义务格按句两条 claim；owner=cell，行 claim 不再携带该格文本
            self.assertEqual(len(cell_claims), 2)
            self.assertEqual(
                {row["locator"]["table_cell_id"] for row in cell_claims},
                {"TBL-000001-R000002-C000002"},
            )
            row_claim = next(
                row for row in catalog["catalog"]
                if row["source_kind"] == "table_row" and "Signing" in row["text"]
            )
            # 代词主体（P1-4）："It shall sign responses." 句内无具体主体，
            # 行头身份必须随 claim 文本下发（B1 结构化同行身份格）
            self.assertIn("Feature=Signing", row_claim["text"])
            self.assertIn("It shall sign responses.", row_claim["text"])
            target_claim = next(
                row for row in cell_claims if "shall authenticate" in row["text"]
            )
            # F4：cell claim 带确定性 semantic_context（表标题+列头+正文句）
            semantic_context = str(target_claim["semantic_context"])
            self.assertIn("Behavior", semantic_context)
            self.assertTrue(
                semantic_context.rstrip().endswith(
                    "The meter shall authenticate all clients."
                )
            )

            # ---- 3) publish（F5：cell 产物哈希绑定进 generation_meta） ---------
            atomic_write_jsonl(out / "ai_requirements.jsonl", [])
            # 真实 metadata：input_fingerprint 由 blocks.jsonl 现算（refresh 门禁依赖）
            ai_extract.write_ai_requirements_metadata(out)
            shadow = claim_ledger.build_shadow_ledger(catalog, [])
            claim_artifacts.publish_shadow_generation(
                out,
                catalog,
                shadow,
                run_id="e2e-f10",
                requirements_sha256=file_sha256(out / "ai_requirements.jsonl"),
            )
            generation = json.loads(
                (out / claim_artifacts.CLAIM_GENERATION_META).read_text(encoding="utf-8")
            )
            self.assertEqual(
                generation.get("table_cell_items_file_sha256"),
                file_sha256(out / "table_cell_items.jsonl"),
            )

            # ---- 4) fold + queue（table_cell proposal 且过 v3 schema） --------
            fold = claim_review_actions.fold_effective_ledger(out, actor_trigger="e2e-f10")
            self.assertNotEqual(fold.get("error"), "base_migration_required")
            snapshot = claim_artifacts.load_committed_effective_snapshot_readonly(out)
            proposals = snapshot["queue_proposals"]
            cell_proposals = [
                row for row in proposals
                if (row.get("focus") or {}).get("kind") == "table_cell"
            ]
            self.assertEqual(len(cell_proposals), 2)
            self.assertNotIn(
                excluded_cell_candidate["claim_id"],
                {row["claim_id"] for row in proposals},
            )
            schema = json.loads(
                Path("schemas/claim_queue_proposal_v3.schema.json").read_text(encoding="utf-8")
            )
            validator = Draft202012Validator(schema)
            for proposal in cell_proposals:
                validator.validate(proposal)
            missing_title = copy.deepcopy(cell_proposals[0])
            missing_title["focus"].pop("table_title")
            with self.assertRaises(ValidationError):
                validator.validate(missing_title)
            proposal = next(
                row for row in cell_proposals
                if "shall authenticate" in str((row.get("focus") or {}).get("text") or "")
            )
            self.assertEqual(proposal["focus"]["table_cell_id"],
                             target_claim["locator"]["table_cell_id"])
            self.assertEqual(proposal["focus"]["table_title"], "Table 1")

            # ---- 5) reload + 篡改 fail-closed ----------------------------------
            claim_artifacts.load_committed_claim_base(out)

            # ---- 6) execute（B2 诚实化：唯一 mock 是最低层传输函数 _post_json）----
            # critique_section/质量刷新/metadata/compliance/merged_spec/freshness/
            # reload 全部走真实实现；config 经 env 真实解析。模拟传输复刻
            # _post_json 的预算 reserve/commit 与 stats 契约，响应内容一律从
            # 真实请求 prompt 派生——prompt 缺对象身份/列头/逐字证据时 claim 无法闭合。
            preconditions = proposal["execution_preconditions"]
            recorded_payloads: list[dict[str, Any]] = []

            with mock.patch.dict(os.environ, _ROUTE_ENV, clear=False), mock.patch.object(
                llm_client, "_post_json",
                side_effect=_make_fake_post_json(recorded_payloads),
            ):
                # 与 execute 同一 env 下走真实 config 解析（env 覆盖 → yaml route →
                # apply_min_tokens 下限），preflight/execute 两侧真实同源。
                route_revision = execution.claim_queue_route_preflight(
                    "openai_compatible"
                )["route_config_revision"]
                result = execution.execute_claim_queue_proposal(
                    out,
                    proposal_id=proposal["proposal_id"],
                    expected_claim_effective_revision=(
                        preconditions["expected_claim_effective_revision"]
                    ),
                    expected_ledger_state=preconditions["expected_ledger_state"],
                    actor="expert:e2e",
                    allow_llm=True,
                    route="openai_compatible",
                    maximum_calls=4,
                    # 真实 apply_min_tokens 把 extract max_tokens 抬到 6144 下限——
                    # 单次调用的 reserve ceiling ≈ body + max_tokens，预算按真实
                    # 下限给足；committed usage 仍按传输层真实回报核算。
                    total_token_budget=200000,
                    request_idempotency_key="e2e-f10-request-1",
                    expected_route_config_revision=route_revision,
                )

            self.assertEqual(result["lifecycle"], "executed")
            self.assertEqual(result["resolution"], "covered")
            # 崩溃 usage 窗口（P1-4/P0-1）：终态 usage 与传输层付费调用逐次相等——
            # 少记一次都说明恢复数据源失真（进程级崩溃探针见 P0-1 专项）
            self.assertEqual(result["usage"]["calls"], len(recorded_payloads))
            self.assertEqual(
                result["usage"]["total_tokens"], 96 * len(recorded_payloads)
            )

            # ---- 7) 真实 prompt 取证（B2 核心断言）--------------------------
            # 对象身份（B1 行头上下文）+ 列头 + 逐字义务句必须真实进入付费 prompt；
            # 模拟模型的产出全部由该 prompt 派生，任一缺失 claim 都不可能闭合。
            extraction_prompts = [
                str(message.get("content") or "")
                for payload in recorded_payloads
                for message in (payload.get("messages") or [])
                if message.get("role") == "user"
                and "【严格定向需求抽取】" in str(message.get("content") or "")
            ]
            self.assertEqual(len(extraction_prompts), 1)
            extraction_prompt = extraction_prompts[0]
            # P0-3 结构化证据段：逐字句在可引用证据，行头/列头在定位上下文
            self.assertIn("可引用证据（", extraction_prompt)
            self.assertIn("定位上下文（", extraction_prompt)
            self.assertNotIn("唯一允许的 focus evidence", extraction_prompt)
            self.assertIn("Encryption", extraction_prompt)
            self.assertIn("Behavior", extraction_prompt)
            self.assertIn(
                "The meter shall authenticate all clients.", extraction_prompt
            )
            _assert_paid_surface_closed(self, recorded_payloads)

            # ---- 8) 闭环后重载：目标 claim 以 covered 落账，其余 cell claim 不受影响
            reloaded = claim_artifacts.load_committed_effective_snapshot_readonly(out)
            ledger_rows = {
                row["claim_id"]: row for row in reloaded["effective_ledger"]
            }
            self.assertEqual(
                ledger_rows[target_claim["claim_id"]]["resolution"], "covered"
            )
            other = next(
                row for row in cell_claims
                if row["claim_id"] != target_claim["claim_id"]
            )
            self.assertNotEqual(
                ledger_rows[other["claim_id"]]["resolution"], "covered"
            )

            # ---- 9) annotation 投影（P0-2 真实链路）：cell claim 必须进入批注
            rendered = dae.render_annotation_html(out, layout_mode="optimized")
            claims_payload = json.loads(
                re.search(r"const CLAIMS = (\[.*?\]);\n", rendered).group(1)
            )
            projected_cell_ids = {
                row["claim_id"] for row in claims_payload
                if row["source_kind"] == "table_cell"
                and row["eligibility"] == "claim"
            }
            self.assertEqual(
                projected_cell_ids, {row["claim_id"] for row in cell_claims}
            )
            projected_candidate = next(
                row for row in claims_payload
                if row["claim_id"] == excluded_cell_candidate["claim_id"]
            )
            self.assertEqual(projected_candidate["eligibility"], "excluded")
            self.assertEqual(projected_candidate["classification"], "non_normative")

    def test_marker_matrix_cell_claim_closes_with_composite_binding(self) -> None:
        """真正的 mapping_matrix（P1-4）：marker 格 claim 全链闭合。

        复合事实（主体×维度×marker）必须三者同现：模拟抽取从「矩阵事实」段
        派生、模拟验证器逐项核对行头主体与列头维度，缺一则 covered=false。
        """
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx_path = root / "marker.docx"
            _write_marker_matrix_docx(docx_path)
            out = root / "out"
            out.mkdir()

            blocks, items, cells = extract_docx(docx_path)
            atomic_write_jsonl(out / "blocks.jsonl", blocks)
            atomic_write_jsonl(out / "table_items.jsonl", items)
            atomic_write_jsonl(out / "table_cell_items.jsonl", cells)

            table_block = next(
                block for block in blocks if block.get("type") == "table"
            )
            self.assertEqual(table_block.get("table_kind"), "mapping_matrix")

            catalog = claim_catalog.build_claim_catalog(
                blocks, items, table_cell_items=cells
            )
            self.assertEqual(catalog["meta"]["table_structure_status"], "ok")
            cell_claims = [
                row for row in catalog["catalog"]
                if row["source_kind"] == "table_cell"
            ]
            # 3 个 marker 格 → 3 条 cell claim（Data access×GET/SET + Event push×GET）
            self.assertEqual(len(cell_claims), 3)
            target_claim = next(
                row for row in cell_claims
                if row["locator"]["table_cell_id"] == "TBL-000001-R000002-C000002"
            )
            self.assertEqual(
                str(target_claim["semantic_context"]),
                "Table 1 - Optical communication interface capabilities | "
                "Interface=Data access | GET = X",
            )

            atomic_write_jsonl(out / "ai_requirements.jsonl", [])
            ai_extract.write_ai_requirements_metadata(out)
            shadow = claim_ledger.build_shadow_ledger(catalog, [])
            claim_artifacts.publish_shadow_generation(
                out,
                catalog,
                shadow,
                run_id="e2e-f10-marker",
                requirements_sha256=file_sha256(out / "ai_requirements.jsonl"),
            )
            claim_review_actions.fold_effective_ledger(
                out, actor_trigger="e2e-f10-marker"
            )
            snapshot = claim_artifacts.load_committed_effective_snapshot_readonly(out)
            proposal = next(
                row for row in snapshot["queue_proposals"]
                if (row.get("focus") or {}).get("table_cell_id")
                == "TBL-000001-R000002-C000002"
            )
            self.assertEqual(
                proposal["focus"]["table_title"],
                "Table 1 - Optical communication interface capabilities",
            )
            preconditions = proposal["execution_preconditions"]
            recorded_payloads: list[dict[str, Any]] = []

            with mock.patch.dict(os.environ, _ROUTE_ENV, clear=False), mock.patch.object(
                llm_client, "_post_json",
                side_effect=_make_fake_post_json(recorded_payloads),
            ):
                route_revision = execution.claim_queue_route_preflight(
                    "openai_compatible"
                )["route_config_revision"]
                result = execution.execute_claim_queue_proposal(
                    out,
                    proposal_id=proposal["proposal_id"],
                    expected_claim_effective_revision=(
                        preconditions["expected_claim_effective_revision"]
                    ),
                    expected_ledger_state=preconditions["expected_ledger_state"],
                    actor="expert:e2e",
                    allow_llm=True,
                    route="openai_compatible",
                    maximum_calls=4,
                    total_token_budget=200000,
                    request_idempotency_key="e2e-f10-marker-request-1",
                    expected_route_config_revision=route_revision,
                )

            self.assertEqual(result["lifecycle"], "executed")
            self.assertEqual(result["resolution"], "covered")
            self.assertEqual(result["usage"]["calls"], len(recorded_payloads))

            # 复合事实段真实进入付费 prompt（主体/维度/取值三者齐发）
            extraction_prompts = [
                str(message.get("content") or "")
                for payload in recorded_payloads
                for message in (payload.get("messages") or [])
                if message.get("role") == "user"
                and "【严格定向需求抽取】" in str(message.get("content") or "")
            ]
            self.assertEqual(len(extraction_prompts), 1)
            extraction_prompt = extraction_prompts[0]
            self.assertIn(
                "表标题：Table 1 - Optical communication interface capabilities",
                extraction_prompt,
            )
            self.assertIn("矩阵事实：主体=Interface=Data access", extraction_prompt)
            self.assertIn("维度=GET", extraction_prompt)
            self.assertIn("取值=X", extraction_prompt)
            _assert_paid_surface_closed(self, recorded_payloads)

            published_requirements = ai_extract.read_jsonl(
                out / ai_extract.AI_REQUIREMENTS
            )
            self.assertTrue(
                any(
                    "Optical communication interface capabilities"
                    in str(row.get("description") or "")
                    and str(row.get("source_quote") or "") == "X"
                    for row in published_requirements
                )
            )

            # 宽候选 + 严闭合（_candidate_basis 格全文精确绑定豁免 6-alnum 下限）：
            # 三条 marker claim 全部到达独立 verifier，由七维按完整
            # semantic_context 逐条裁定——不是只放行了目标那一条
            verifier_requests = [
                json.loads(str(message.get("content") or ""))
                for payload in recorded_payloads
                for message in (payload.get("messages") or [])
                if message.get("role") == "user"
                and "claim-coverage-verifier-request/v2"
                in str(message.get("content") or "")
            ]
            self.assertEqual(len(verifier_requests), 1)
            verified_contexts = {
                str(group[1])
                for group in verifier_requests[0].get("groups") or []
            }
            self.assertTrue(
                {str(row["semantic_context"]) for row in cell_claims}.issubset(
                    verified_contexts
                ),
                "every marker-cell claim must reach the independent verifier",
            )

            # 闭环后重载：目标 marker claim covered，同表其余 marker claim 不受影响
            reloaded = claim_artifacts.load_committed_effective_snapshot_readonly(out)
            ledger_rows = {
                row["claim_id"]: row for row in reloaded["effective_ledger"]
            }
            self.assertEqual(
                ledger_rows[target_claim["claim_id"]]["resolution"], "covered"
            )
            for row in cell_claims:
                if row["claim_id"] != target_claim["claim_id"]:
                    self.assertNotEqual(
                        ledger_rows[row["claim_id"]]["resolution"], "covered"
                    )
            caption_claim = next(
                row for row in catalog["catalog"]
                if row["text"]
                == "Table 1 - Optical communication interface capabilities"
            )
            self.assertNotEqual(
                ledger_rows[caption_claim["claim_id"]]["resolution"], "covered"
            )

            # annotation 投影（P0-2/P1-3 真实链路）：marker cell claim 按物理
            # R×C 落进静态 HTML 对应 <td>（Data access 行的 GET/SET 格各一枚）
            rendered = dae.render_annotation_html(out, layout_mode="optimized")
            claims_payload = json.loads(
                re.search(r"const CLAIMS = (\[.*?\]);\n", rendered).group(1)
            )
            projected_cell_ids = {
                row["claim_id"] for row in claims_payload
                if row["source_kind"] == "table_cell"
            }
            self.assertEqual(
                projected_cell_ids, {row["claim_id"] for row in cell_claims}
            )
            chip_cells = re.findall(
                r'class="claim-cell-chip [^"]*"[^>]*data-table-cell-id="([^"]+)"',
                rendered,
            )
            self.assertEqual(
                sorted(chip_cells),
                [
                    "TBL-000001-R000002-C000002",
                    "TBL-000001-R000002-C000003",
                    "TBL-000001-R000003-C000002",
                ],
            )
            body_rows = re.findall(r"<tr[^>]*>(.*?)</tr>", rendered, re.DOTALL)
            data_access_row = next(
                row_html for row_html in body_rows if "Data access" in row_html
            )
            tds = re.findall(r"<td>(.*?)</td>", data_access_row, re.DOTALL)
            self.assertNotIn("claim-cell-chip", tds[0])
            self.assertIn("claim-cell-chip", tds[1])
            self.assertIn("claim-cell-chip", tds[2])


class XlsxRegionConservationTests(unittest.TestCase):
    """xlsx sheet 级非空格守恒计数器（F10）：漏格/重叠 = parse_incomplete 硬失败。"""

    @staticmethod
    def _extract(path: Path):
        return extract_xlsx(path, knowledge_bases=KnowledgeRepository.from_paths([]))

    def test_listobject_leftover_region_is_conserved(self) -> None:
        # ListObject 矩形之外的非空连通区域同样成表（表外需求不得消失）
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "book.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Specs"
            sheet["A1"] = "Name"
            sheet["B1"] = "Value"
            sheet["A2"] = "Voltage"
            sheet["B2"] = "230 V"
            sheet.add_table(ExcelTable(displayName="MainTable", ref="A1:B2"))
            sheet["D5"] = "Alarm"
            sheet["E5"] = "The meter shall raise alarms."
            workbook.save(path)

            blocks, items, cells = self._extract(path)

        tables = [block for block in blocks if block["type"] == "table"]
        self.assertEqual(len(tables), 2)
        self.assertTrue(all(not block.get("parse_incomplete") for block in tables))
        flat = " ".join(str(block.get("text") or "") for block in tables)
        self.assertIn("shall raise alarms", flat)

    def test_plain_sheet_has_no_parse_incomplete(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "plain.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Plain"
            sheet.append(["Name", "Requirement"])
            sheet.append(["Logging", "The meter shall log events."])
            workbook.save(path)

            blocks, _items, _cells = self._extract(path)

        tables = [block for block in blocks if block["type"] == "table"]
        self.assertEqual(len(tables), 1)
        self.assertFalse(bool(tables[0].get("parse_incomplete")))

    def test_region_gap_trips_conservation_counter(self) -> None:
        # 区域拆分漏格（模拟 ListObject 之外连通区域丢失）→ 计数器硬失败，
        # 表块标 parse_incomplete 且目录账本如实 incomplete（宁 incomplete 不丢字）
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "gap.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Gap"
            sheet["A1"] = "Name"
            sheet["B1"] = "Value"
            sheet["A2"] = "Voltage"
            sheet["B2"] = "230 V"
            sheet.add_table(ExcelTable(displayName="MainTable", ref="A1:B2"))
            sheet["D5"] = "Alarm"
            sheet["E5"] = "The meter shall raise alarms."
            workbook.save(path)

            import parsers.xlsx_parser as xlsx_module

            with mock.patch.object(
                xlsx_module, "_connected_regions", return_value=[]
            ):
                blocks, _items, cells = self._extract(path)

        tables = [block for block in blocks if block["type"] == "table"]
        self.assertTrue(tables)
        self.assertTrue(all(bool(block.get("parse_incomplete")) for block in tables))
        reason = tables[0].get("parse_incomplete_reason") or {}
        self.assertEqual(reason.get("code"), "xlsx_region_conservation")
        self.assertGreaterEqual(int(reason.get("dropped_cell_count") or 0), 2)
        catalog = claim_catalog.build_claim_catalog(blocks, _items, table_cell_items=cells)
        self.assertEqual(catalog["meta"]["accounting_status"], "incomplete")
        self.assertGreaterEqual(catalog["meta"]["audit"].get("parse_incomplete_count") or 0, 1)


if __name__ == "__main__":
    unittest.main()
