"""P5 导出：把装配好的《DLMS/COSEM 实现规格》渲染成人读交付物（Word + Markdown）。

assemble_spec 产出的 dlms_cosem_spec_requirements.json 是喂公司工具链的机器格式；
本模块在其上叠加给研发直接阅读的文档：按功能域分组、每条带溯源（章节 + 原文）、
属性表/矩阵渲染为真正的表格、附"还需查阅的外部规范"清单。零 LLM、零文档外内容。

用法：python -m spec_export --out <atomizer 输出目录> [--format md,docx] [--reviews <jsonl>]
"""
from __future__ import annotations

import argparse
import datetime
import json
from collections import defaultdict
from pathlib import Path
from typing import Any

# 功能域分组顺序（取自 requirement_schema 的标签体系）；未列出的域追加在「其它」前
DOMAIN_ORDER = ["安全", "计量", "通信协议", "时钟", "事件记录", "费率"]
OTHER_DOMAIN = "其它"
PRIORITY_ORDER = {"P0": 0, "P1": 1, "P2": 2}
ASSEMBLED_JSON = "dlms_cosem_spec_requirements.json"


def load_doc(out_dir: Path, reviews_path: Path | None) -> dict[str, Any]:
    """优先读已装配的 JSON；不存在则现造（单一数据源 = assemble_spec.assemble）。"""
    out_dir = out_dir.expanduser().resolve()
    target = out_dir / ASSEMBLED_JSON
    if target.exists():
        return json.loads(target.read_text(encoding="utf-8"))
    from assemble_spec import assemble  # 延迟导入，避免循环
    doc, _ = assemble(out_dir, reviews_path, source=out_dir.name,
                      extracted_at=datetime.datetime.now().isoformat(timespec="seconds"))
    return doc


def primary_domain(req: dict[str, Any]) -> str:
    labels = req.get("labels") or []
    return str(labels[0]) if labels else OTHER_DOMAIN


def group_by_domain(requirements: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for req in requirements:
        groups[primary_domain(req)].append(req)
    for reqs in groups.values():
        reqs.sort(key=lambda r: (PRIORITY_ORDER.get(str(r.get("priority")), 9), str(r.get("id"))))
    ordered_domains = [d for d in DOMAIN_ORDER if d in groups]
    extras = sorted(d for d in groups if d not in DOMAIN_ORDER and d != OTHER_DOMAIN)
    if OTHER_DOMAIN in groups:
        extras.append(OTHER_DOMAIN)
    return [(d, groups[d]) for d in [*ordered_domains, *extras]]


# --- Markdown -------------------------------------------------------------------

def render_markdown(doc: dict[str, Any]) -> str:
    meta = doc.get("meta", {})
    analysis = doc.get("analysis", {})
    requirements = doc.get("requirements", [])
    lines: list[str] = [
        "# DLMS/COSEM 实现规格",
        "",
        f"- 来源：`{meta.get('source', '')}`",
        f"- 生成时间：{meta.get('extracted_at', '')}",
        f"- 目标标准：{', '.join(meta.get('target_standards', []))}",
        f"- 需求总数：**{analysis.get('total_count', len(requirements))}**",
        "",
        "## 概要",
        "",
        f"- 按优先级：{_kv(analysis.get('by_priority', {}))}",
        f"- 按类型：{_kv(analysis.get('by_type', {}))}",
        f"- {analysis.get('coverage_report', '')}",
        "",
    ]

    for domain, reqs in group_by_domain(requirements):
        lines.append(f"## {domain}（{len(reqs)}）")
        lines.append("")
        for req in reqs:
            lines.extend(_md_requirement(req))

    lines.extend(_md_external_refs(doc.get("external_references")))
    lines.extend(_md_gaps(analysis.get("gaps") or []))
    return "\n".join(lines)


def _kv(mapping: dict[str, Any]) -> str:
    return "　".join(f"{k} {v}" for k, v in mapping.items()) or "—"


def _md_requirement(req: dict[str, Any]) -> list[str]:
    lines = [
        f"### {req.get('id')} {req.get('title')}",
        "",
        f"`{req.get('priority')}` · `{req.get('status')}` · `{req.get('type')}`",
        "",
        str(req.get("description") or ""),
        "",
    ]
    section = req.get("source_section") or ""
    quote = req.get("source_quote") or ""
    if section or quote:
        lines.append(f"**溯源**：{section}")
        if quote:
            lines.append(f"> {quote}")
        lines.append("")
    tt = req.get("threshold_table")
    if tt and tt.get("columns"):
        if tt.get("description"):
            lines.append(f"*{tt['description']}*")
            lines.append("")
        lines.append("| " + " | ".join(str(c) for c in tt["columns"]) + " |")
        lines.append("|" + "|".join("---" for _ in tt["columns"]) + "|")
        for row in tt.get("rows", []):
            cells = [str(c).replace("\n", " ").replace("|", "\\|") for c in row]
            cells += [""] * (len(tt["columns"]) - len(cells))
            lines.append("| " + " | ".join(cells[: len(tt["columns"])]) + " |")
        lines.append("")
    criteria = req.get("acceptance_criteria") or []
    if criteria:
        lines.append("**验收标准**：")
        for c in criteria:
            lines.append(f"- [ ] {c}")
        lines.append("")
    if req.get("notes"):
        lines.append(f"_注：{req['notes']}_")
        lines.append("")
    return lines


def _md_external_refs(ext: dict[str, Any] | None) -> list[str]:
    if not ext or not ext.get("references"):
        return []
    counts = ext.get("counts", {})
    lines = [
        "## 附录 A：外部规范交叉引用（实现还需查阅）",
        "",
        f"> 必查 {counts.get('normative', 0)} / 仅列目录 {counts.get('listed_only', 0)}；"
        "确定性扫描，研发实现协议/行为层时按此查阅外部标准。",
        "",
    ]
    normative = [e for e in ext["references"] if e.get("materiality") == "normative"]
    listed = [e for e in ext["references"] if e.get("materiality") != "normative"]
    if normative:
        lines.append("### 必查")
        lines.append("")
        for e in normative:
            note = f"　〔{e['dlms_ua_note']}〕" if e.get("dlms_ua_note") else ""
            lines.append(f"- **{e['spec_id']}** — {e.get('title') or '（文档未给标题）'}{note}")
        lines.append("")
    if listed:
        lines.append("### 仅出现在引用目录")
        lines.append("")
        for e in listed:
            lines.append(f"- {e['spec_id']} — {e.get('title') or '（文档未给标题）'}")
        lines.append("")
    return lines


def _md_gaps(gaps: list[dict[str, Any]]) -> list[str]:
    if not gaps:
        return []
    lines = ["## 附录 B：覆盖缺口", "", "> 以下适配域未发现相关需求，请核对原文是否遗漏。", ""]
    for g in gaps:
        lines.append(f"- **{g.get('domain')}**：{g.get('description')}")
    lines.append("")
    return lines


# --- Word -----------------------------------------------------------------------

def write_docx(doc: dict[str, Any], path: Path) -> None:
    from docx import Document  # 已有依赖（输入解析在用）

    meta = doc.get("meta", {})
    analysis = doc.get("analysis", {})
    requirements = doc.get("requirements", [])

    document = Document()
    document.add_heading("DLMS/COSEM 实现规格", level=0)
    document.add_paragraph(f"来源：{meta.get('source', '')}")
    document.add_paragraph(f"生成时间：{meta.get('extracted_at', '')}")
    document.add_paragraph(f"目标标准：{', '.join(meta.get('target_standards', []))}")
    document.add_paragraph(f"需求总数：{analysis.get('total_count', len(requirements))}")

    document.add_heading("概要", level=1)
    document.add_paragraph(f"按优先级：{_kv(analysis.get('by_priority', {}))}")
    document.add_paragraph(f"按类型：{_kv(analysis.get('by_type', {}))}")
    document.add_paragraph(str(analysis.get("coverage_report", "")))

    for domain, reqs in group_by_domain(requirements):
        document.add_heading(f"{domain}（{len(reqs)}）", level=1)
        for req in reqs:
            _docx_requirement(document, req)

    _docx_external_refs(document, doc.get("external_references"))
    _docx_gaps(document, analysis.get("gaps") or [])
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _docx_requirement(document: Any, req: dict[str, Any]) -> None:
    document.add_heading(f"{req.get('id')} {req.get('title')}", level=2)
    document.add_paragraph(f"{req.get('priority')} · {req.get('status')} · {req.get('type')}")
    document.add_paragraph(str(req.get("description") or ""))
    section = req.get("source_section") or ""
    quote = req.get("source_quote") or ""
    if section or quote:
        document.add_paragraph(f"溯源：{section}")
        if quote:
            document.add_paragraph(str(quote)).style = "Quote"
    tt = req.get("threshold_table")
    if tt and tt.get("columns"):
        if tt.get("description"):
            document.add_paragraph(str(tt["description"]))
        columns = [str(c) for c in tt["columns"]]
        table = document.add_table(rows=1, cols=len(columns))
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        for i, col in enumerate(columns):
            table.rows[0].cells[i].text = col
        for row in tt.get("rows", []):
            cells = table.add_row().cells
            for i in range(len(columns)):
                cells[i].text = str(row[i]) if i < len(row) else ""
    criteria = req.get("acceptance_criteria") or []
    if criteria:
        document.add_paragraph("验收标准：")
        for c in criteria:
            document.add_paragraph(str(c), style="List Bullet")
    if req.get("notes"):
        document.add_paragraph(f"注：{req['notes']}")


def _docx_external_refs(document: Any, ext: dict[str, Any] | None) -> None:
    if not ext or not ext.get("references"):
        return
    counts = ext.get("counts", {})
    document.add_heading("附录 A：外部规范交叉引用（实现还需查阅）", level=1)
    document.add_paragraph(
        f"必查 {counts.get('normative', 0)} / 仅列目录 {counts.get('listed_only', 0)}；"
        "确定性扫描，研发实现协议/行为层时按此查阅外部标准。")
    normative = [e for e in ext["references"] if e.get("materiality") == "normative"]
    listed = [e for e in ext["references"] if e.get("materiality") != "normative"]
    if normative:
        document.add_heading("必查", level=2)
        for e in normative:
            note = f"　〔{e['dlms_ua_note']}〕" if e.get("dlms_ua_note") else ""
            document.add_paragraph(f"{e['spec_id']} — {e.get('title') or '（文档未给标题）'}{note}",
                                   style="List Bullet")
    if listed:
        document.add_heading("仅出现在引用目录", level=2)
        for e in listed:
            document.add_paragraph(f"{e['spec_id']} — {e.get('title') or '（文档未给标题）'}",
                                   style="List Bullet")


def _docx_gaps(document: Any, gaps: list[dict[str, Any]]) -> None:
    if not gaps:
        return
    document.add_heading("附录 B：覆盖缺口", level=1)
    document.add_paragraph("以下适配域未发现相关需求，请核对原文是否遗漏。")
    for g in gaps:
        document.add_paragraph(f"{g.get('domain')}：{g.get('description')}", style="List Bullet")


# --- 入口 -----------------------------------------------------------------------

def export_spec(out_dir: Path, *, formats: list[str], reviews_path: Path | None = None) -> list[str]:
    out_dir = out_dir.expanduser().resolve()
    doc = load_doc(out_dir, reviews_path)
    written: list[str] = []
    wanted = {f.strip().lower() for f in formats if f.strip()}
    if "md" in wanted:
        (out_dir / "dlms_cosem_spec.md").write_text(render_markdown(doc), encoding="utf-8")
        written.append("dlms_cosem_spec.md")
    if "docx" in wanted:
        write_docx(doc, out_dir / "dlms_cosem_spec.docx")
        written.append("dlms_cosem_spec.docx")
    if "xlsx" in wanted:
        from spec_excel import write_xlsx

        write_xlsx(doc, out_dir / "dlms_cosem_spec.xlsx")
        written.append("dlms_cosem_spec.xlsx")
    unknown = wanted - {"md", "docx", "xlsx"}
    if unknown:
        raise ValueError(f"unsupported export format: {', '.join(sorted(unknown))}")
    return written


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render the assembled DLMS/COSEM spec to Word + Markdown.")
    parser.add_argument("--out", type=Path, required=True, help="Atomizer output directory")
    parser.add_argument("--format", default="md,docx", help="Comma-separated: md,docx")
    parser.add_argument("--reviews", type=Path, default=None, help="Behaviour LLM reviews jsonl (used if assembling fresh)")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    written = export_spec(args.out, formats=args.format.split(","), reviews_path=args.reviews)
    print(json.dumps({"out": str(args.out.expanduser().resolve()), "written": written},
                     ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
