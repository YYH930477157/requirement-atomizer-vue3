from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

import spec_export


def write_jsonl(path: Path, rows: list[dict]) -> None:
    path.write_text("\n".join(json.dumps(r, ensure_ascii=False) for r in rows) + "\n", encoding="utf-8")


def write_fixture(out_dir: Path) -> None:
    write_jsonl(out_dir / "atomic_requirements.jsonl", [
        {"stable_req_id": "O-CLOCK", "requirement_type": "cosem_object_instance",
         "object": "Clock", "domain": "time", "source_refs": ["TBL-1"], "section_path": ["Time"]},
        {"stable_req_id": "A-CLOCK", "requirement_type": "cosem_attribute_access",
         "object": "Clock.time", "source_refs": ["TBL-2"], "verification_method": "configuration_check"},
        {"stable_req_id": "ASM-1", "requirement_type": "association_security_matrix",
         "object": "ClientA", "source_refs": ["TBL-3"]},
        {"stable_req_id": "F-1", "requirement_type": "functional", "object": "Counter",
         "requirement": "The meter must reset to 0 per IEC 62056-5-3.", "source_refs": ["BLK-1"],
         "section_path": ["Behaviour"]},
    ])
    write_jsonl(out_dir / "table_items.jsonl", [
        {"item_id": "TBL-1", "fields": {"Object/attribute name": "Clock", "CL": "8", "Value": "0-0:1.0.0.255"}},
        {"item_id": "TBL-2", "fields": {"#": "1", "Object/attribute name": "time", "Type": "octet-string",
                                         "Value": "00", "Access rights RC/PC/SC/LC": "R-/RW/--/R-"}},
        {"item_id": "TBL-3", "fields": {"Customer application process": "ClientA",
                                         "Server application process / Management Logical Device": "HLS"}},
    ])
    write_jsonl(out_dir / "blocks.jsonl", [
        {"block_id": "BLK-1", "type": "paragraph", "section_path": ["Behaviour"],
         "text": "The counter behaviour follows IEC 62056-5-3."},
        {"block_id": "BLK-2", "type": "heading", "section_path": ["Normative references"],
         "text": "IEC 62056-5-3, Electricity metering - Part 5-3: COSEM application layer"},
    ])


class SpecExportTests(unittest.TestCase):
    def test_exports_md_and_docx(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            write_fixture(out)
            written = spec_export.export_spec(out, formats=["md", "docx"])
            self.assertEqual(set(written), {"dlms_cosem_spec.md", "dlms_cosem_spec.docx"})
            self.assertTrue((out / "dlms_cosem_spec.md").exists())
            self.assertTrue((out / "dlms_cosem_spec.docx").exists())

    def test_markdown_structure(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            write_fixture(out)
            spec_export.export_spec(out, formats=["md"])
            md = (out / "dlms_cosem_spec.md").read_text(encoding="utf-8")

        self.assertIn("# DLMS/COSEM 实现规格", md)
        self.assertRegex(md, r"## .+（\d+）")                     # 功能域分组标题带计数
        self.assertRegex(md, r"### REQ-\d{3} ")                    # 需求条目
        self.assertIn("**溯源**", md)                              # 溯源段
        self.assertIn("| # | 属性 | 类型 |", md)                   # threshold 表渲染
        self.assertIn("附录 A：外部规范交叉引用", md)              # P4 附录
        self.assertIn("IEC 62056-5-3", md)                        # 外部规范被列出

    def test_docx_opens_with_headings_and_tables(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            write_fixture(out)
            spec_export.export_spec(out, formats=["docx"])
            document = Document(str(out / "dlms_cosem_spec.docx"))

        self.assertGreater(len(document.paragraphs), 0)
        self.assertGreater(len(document.tables), 0)               # 至少一个属性/矩阵表
        self.assertTrue(any("DLMS/COSEM 实现规格" in p.text for p in document.paragraphs))

    def test_unknown_format_raises(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            write_fixture(out)
            with self.assertRaises(ValueError):
                spec_export.export_spec(out, formats=["pdf"])

    def test_reuses_existing_assembled_json(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            # 预置一份最小装配 JSON，断言导出直接复用它（不依赖原子文件）
            (out / "dlms_cosem_spec_requirements.json").write_text(json.dumps({
                "meta": {"source": "preset", "extracted_at": "t", "target_standards": ["X"]},
                "requirements": [{"id": "REQ-001", "title": "T", "description": "D", "type": "functional",
                                  "priority": "P1", "status": "draft", "source_section": "S",
                                  "source_quote": "Q", "threshold_table": None,
                                  "acceptance_criteria": ["c1"], "labels": ["安全"], "notes": ""}],
                "analysis": {"total_count": 1, "by_priority": {"P1": 1}, "by_type": {"functional": 1},
                             "coverage_report": "rpt", "gaps": []},
            }, ensure_ascii=False), encoding="utf-8")
            spec_export.export_spec(out, formats=["md"])
            md = (out / "dlms_cosem_spec.md").read_text(encoding="utf-8")
        self.assertIn("preset", md)
        self.assertIn("## 安全（1）", md)

    def test_xlsx_includes_object_model_sheet_in_obis_list_format(self) -> None:
        from openpyxl import load_workbook

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            write_fixture(out)
            spec_export.export_spec(out, formats=["xlsx"])
            workbook = load_workbook(out / "dlms_cosem_spec.xlsx")

        self.assertIn("COSEM Object Model", workbook.sheetnames)
        ws = workbook["COSEM Object Model"]
        self.assertEqual(
            [ws.cell(row=1, column=col).value for col in range(1, 12)],
            [
                "Index",
                "Object / Attribute Name",
                "Attribute Type",
                "Class",
                "Ver.",
                "SN",
                "OBIS Code / Default Value",
                "Public(16)",
                "Data Readout(3)",
                "Remote Management(1)",
                "Local Management(2)",
            ],
        )
        self.assertEqual(ws["B2"].value, "Time")
        self.assertEqual(ws["B2"].fill.fgColor.rgb, "FFFF0000")
        self.assertEqual(ws["B3"].value, "Clock")
        self.assertEqual(ws["D3"].value, "8")
        self.assertEqual(ws["G3"].value, "0-0:1.0.0.255")
        self.assertEqual(ws["B3"].fill.fgColor.rgb, "FFFFFF00")
        self.assertEqual(ws["A4"].value, "1")
        self.assertEqual(ws["B4"].value, "time")
        self.assertEqual(ws["C4"].value, "octet-string")
        self.assertEqual(ws["G4"].value, "00")
        self.assertEqual(ws["H4"].value, "RW")
        self.assertEqual(ws["I4"].value, "R-")
        self.assertEqual(ws["J4"].value, "--")
        self.assertEqual(ws["K4"].value, "R-")
        self.assertEqual(ws["B4"].fill.fgColor.rgb, "FF00FFFF")


class SpecExcelDedupTests(unittest.TestCase):
    """对象定义型需求只进「COSEM Object Model」sheet，不在领域 sheet 重复出现。"""

    def _doc(self) -> dict:
        return {
            "meta": {"source": "t", "extracted_at": "2026-01-01T00:00:00", "meter_type": "electric",
                     "target_standards": []},
            "requirements": [
                {  # 对象定义型（匹配 _is_object_model_requirement）+ 领域标签 时钟
                    "id": "REQ-001",
                    "title": "Clock (OBIS 0-0:1.0.0.255 / CL 8)",
                    "description": "实现 COSEM 对象 Clock",
                    "type": "functional", "priority": "P1", "status": "confirmed",
                    "source_section": "Time",
                    "source_quote": "COSEM object Clock / CL 8 / OBIS 0-0:1.0.0.255",
                    "threshold_table": {"description": "Clock 属性访问表",
                                         "columns": ["#", "名称", "类型", "RC", "PC", "SC", "LC", "默认值"],
                                         "rows": [["1", "logical_name", "octet-string[6]", "R-", "R-", "R-", "R-", "0000010000FF"]]},
                    "acceptance_criteria": [], "dependencies": [], "parent": None, "children": [],
                    "labels": ["时钟"], "notes": "",
                },
                {  # 同域行为型需求——应留在领域 sheet
                    "id": "REQ-002",
                    "title": "时钟同步行为",
                    "description": "时钟应每日与基准同步",
                    "type": "functional", "priority": "P1", "status": "confirmed",
                    "source_section": "Time", "source_quote": "the clock shall synchronise daily",
                    "threshold_table": None,
                    "acceptance_criteria": [], "dependencies": [], "parent": None, "children": [],
                    "labels": ["时钟"], "notes": "",
                },
            ],
            "analysis": {"total_count": 2, "by_type": {}, "by_priority": {}, "by_domain": {},
                         "gaps": [], "conflicts": []},
        }

    def test_object_requirement_only_in_object_model_sheet(self) -> None:
        import spec_excel
        from openpyxl import load_workbook

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "spec.xlsx"
            spec_excel.write_xlsx(self._doc(), out)
            wb = load_workbook(out)

        def sheet_contains(name: str, needle: str) -> bool:
            ws = wb[name]
            return any(needle in str(c) for row in ws.iter_rows(values_only=True) for c in row if c is not None)

        self.assertIn("COSEM Object Model", wb.sheetnames)
        self.assertIn("时钟", wb.sheetnames)
        # 对象 OBIS 在 Object Model sheet 在、在领域 sheet 不在
        self.assertTrue(sheet_contains("COSEM Object Model", "0-0:1.0.0.255"))
        self.assertFalse(sheet_contains("时钟", "0-0:1.0.0.255"))
        self.assertFalse(sheet_contains("时钟", "REQ-001"))
        # 行为型需求仍在领域 sheet
        self.assertTrue(sheet_contains("时钟", "REQ-002"))


if __name__ == "__main__":
    unittest.main()
