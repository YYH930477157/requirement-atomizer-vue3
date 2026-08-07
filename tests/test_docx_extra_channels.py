"""Tests for parsers/docx_extra_channels.py (A5②)."""
from __future__ import annotations

import json
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from atomize import extract_docx
import atomize as _atomize_module
from parsers.docx_extra_channels import (
    extract_docx_extra_channels,
    extract_header_footer_texts,
    extract_textbox_texts,
)


def _inject_textbox(doc: Document, text: str) -> None:
    """在 body 第一个段落处注入一个 w:txbxContent 文本框。"""
    body = doc.element.body
    p = body.find(qn("w:p"))
    if p is None:
        p = doc.add_paragraph()._p
    from lxml import etree
    pict = etree.Element(qn("w:pict"))
    shape = etree.SubElement(
        pict,
        "{urn:schemas-microsoft-com:vml}shape",
        {"type": "#_x0000_t202", "style": "position:absolute;"},
    )
    textbox = etree.SubElement(shape, "{urn:schemas-microsoft-com:vml}textbox")
    txbx_content = etree.SubElement(textbox, qn("w:txbxContent"))
    txbx_p = etree.SubElement(txbx_content, qn("w:p"))
    txbx_r = etree.SubElement(txbx_p, qn("w:r"))
    txbx_t = etree.SubElement(txbx_r, qn("w:t"))
    txbx_t.text = text
    r = etree.SubElement(p, qn("w:r"))
    r.append(pict)


class HeaderFooterTests(unittest.TestCase):
    def _make_doc_with_header_footer(self):
        doc = Document()
        doc.add_paragraph("Body paragraph")
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "Header line one"
        header.add_paragraph("Header line two")
        footer = section.footer
        footer.paragraphs[0].text = "Footer line one"
        return doc

    def test_extracts_header_and_footer(self):
        doc = self._make_doc_with_header_footer()
        result = extract_header_footer_texts(doc)
        self.assertIn("Header line one", result["header"])
        self.assertIn("Header line two", result["header"])
        self.assertIn("Footer line one", result["footer"])

    def test_deduplicates_reused_parts(self):
        doc = self._make_doc_with_header_footer()
        result = extract_header_footer_texts(doc)
        self.assertEqual(len(result["header"]), 2)
        self.assertEqual(len(result["footer"]), 1)


class TextboxTests(unittest.TestCase):
    def test_extracts_textbox_text(self):
        doc = Document()
        doc.add_paragraph("Body")
        _inject_textbox(doc, "Textbox content")
        texts = extract_textbox_texts(doc)
        self.assertIn("Textbox content", texts)

    def test_empty_doc_returns_empty(self):
        doc = Document()
        self.assertEqual(extract_textbox_texts(doc), [])


class UnifiedChannelTests(unittest.TestCase):
    def test_channels_are_present(self):
        doc = Document()
        doc.add_paragraph("Body")
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Header text"
        section.footer.paragraphs[0].text = "Footer text"
        _inject_textbox(doc, "Textbox text")
        result = extract_docx_extra_channels(doc)
        self.assertIn("Header text", result["header"])
        self.assertIn("Footer text", result["footer"])
        self.assertIn("Textbox text", result["textbox"])


class ExtractDocxSwitchTests(unittest.TestCase):
    """A5② 开关集成：默认关闭保证 golden blocks.jsonl 不漂移；开启才带 content_channel。"""

    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        os.environ.pop("RATOMIZER_DOCX_EXTRA_CHANNELS", None)

    def tearDown(self):
        self.tmpdir.cleanup()
        os.environ.pop("RATOMIZER_DOCX_EXTRA_CHANNELS", None)

    def _make_docx(self, path: Path) -> None:
        doc = Document()
        doc.add_paragraph("Body paragraph")
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Header line"
        section.footer.paragraphs[0].text = "Footer line"
        _inject_textbox(doc, "Textbox note")
        doc.save(path)

    def _blocks_jsonl(self, blocks: list[dict]) -> str:
        return "\n".join(json.dumps(b, ensure_ascii=False) for b in blocks)

    def test_switch_off_matches_body_only_baseline(self):
        """关闭时 blocks.jsonl 与未启用额外通道的基线逐字节一致。"""
        path = Path(self.tmpdir.name) / "channels.docx"
        self._make_docx(path)

        # 关闭开关的实际输出
        blocks_off, _, _ = extract_docx(path, knowledge_bases=[])
        # 模拟"未引入该功能"的基线：强制开启但让 extra_channels 为空
        os.environ["RATOMIZER_DOCX_EXTRA_CHANNELS"] = "1"
        original = _atomize_module.extract_docx_extra_channels
        try:
            _atomize_module.extract_docx_extra_channels = lambda _doc: {"textbox": [], "header": [], "footer": []}
            blocks_baseline, _, _ = extract_docx(path, knowledge_bases=[])
        finally:
            _atomize_module.extract_docx_extra_channels = original

        self.assertEqual(
            self._blocks_jsonl(blocks_off),
            self._blocks_jsonl(blocks_baseline),
            "关闭开关的输出应与正文-only 基线逐字节一致",
        )
        extra_channels = {b.get("content_channel") for b in blocks_off if b.get("content_channel")}
        self.assertEqual(extra_channels, set())

    def test_switch_on_marks_blocks_with_content_channel(self):
        """开启时 header/footer/textbox 块均带 content_channel 标记。"""
        path = Path(self.tmpdir.name) / "channels.docx"
        self._make_docx(path)
        os.environ["RATOMIZER_DOCX_EXTRA_CHANNELS"] = "1"
        blocks, _, _ = extract_docx(path, knowledge_bases=[])
        channel_texts = {
            str(b.get("content_channel") or "body"): str(b.get("text") or "")
            for b in blocks
        }
        self.assertIn("Header line", channel_texts.get("header", ""))
        self.assertIn("Footer line", channel_texts.get("footer", ""))
        self.assertIn("Textbox note", channel_texts.get("textbox", ""))
        self.assertEqual(channel_texts.get("body"), "Body paragraph")


if __name__ == "__main__":
    unittest.main()
