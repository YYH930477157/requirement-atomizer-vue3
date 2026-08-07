"""A7/A5/A8 集成回归测试（默认开启项）。"""
from __future__ import annotations

import json
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from atomize import run_atomizer_pipeline


from xlsx_requirement_list import BASE_LIBRARY_CANDIDATES_FILE, XLSX_REQUIREMENT_LIST_SWITCH


from claim_quality_rescan import CLAIM_RESCAN_SWITCH


class ClaimRescanIntegrationTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        os.environ.pop(CLAIM_RESCAN_SWITCH, None)

    def tearDown(self):
        self.tmpdir.cleanup()
        os.environ.pop(CLAIM_RESCAN_SWITCH, None)

    def test_claim_rescan_off_by_default(self):
        path = Path(self.tmpdir.name) / "doc.docx"
        doc = Document()
        doc.add_paragraph("Body")
        doc.save(path)
        out_dir = Path(self.tmpdir.name) / "out"
        run_atomizer_pipeline(path, out_dir)
        qp = json.loads((out_dir / "quality_report.json").read_text(encoding="utf-8"))
        self.assertIsNone(qp.get("claim_rescan"))

    def test_claim_rescan_on_reports_issues(self):
        os.environ[CLAIM_RESCAN_SWITCH] = "1"
        path = Path(self.tmpdir.name) / "doc2.docx"
        doc = Document()
        doc.add_paragraph("Body 2")
        doc.save(path)
        out_dir = Path(self.tmpdir.name) / "out2"
        run_atomizer_pipeline(path, out_dir)
        qp = json.loads((out_dir / "quality_report.json").read_text(encoding="utf-8"))
        self.assertIsNotNone(qp.get("claim_rescan"))
        # 纯正文文档通常不会产生 functional_requirements.json，因此会触发 coverage_no_functional_items
        self.assertTrue(qp["claim_rescan"]["enabled"])


class UnextractedRegistryIntegrationTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        os.environ.pop("RATOMIZER_UNEXTRACTED_REGISTRY", None)

    def tearDown(self):
        self.tmpdir.cleanup()

    def _make_docx(self, text: str) -> Path:
        path = Path(self.tmpdir.name) / "doc.docx"
        doc = Document()
        doc.add_paragraph(text)
        doc.save(path)
        return path

    def test_pipeline_writes_unextracted_registry(self):
        path = self._make_docx("Body text")
        out_dir = Path(self.tmpdir.name) / "out"
        manifest = run_atomizer_pipeline(path, out_dir)
        self.assertIn("unextracted_entries", manifest["counts"])
        self.assertTrue((out_dir / "unextracted_registry.json").exists())
        payload = json.loads((out_dir / "unextracted_registry.json").read_text(encoding="utf-8"))
        self.assertEqual(payload["schema"], "unextracted-registry/v1")
        # 纯正文文档无噪声/隐藏 sheet，登记册可为 0
        self.assertIn("total", payload)

    def test_quality_report_has_unextracted_summary(self):
        path = self._make_docx("Body text")
        out_dir = Path(self.tmpdir.name) / "out2"
        run_atomizer_pipeline(path, out_dir)
        qp = json.loads((out_dir / "quality_report.json").read_text(encoding="utf-8"))
        self.assertIn("unextracted_registry", qp)
        self.assertTrue(qp["unextracted_registry"]["available"])


class DocxExtraChannelIntegrationTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        os.environ["RATOMIZER_DOCX_EXTRA_CHANNELS"] = "1"

    def tearDown(self):
        self.tmpdir.cleanup()
        os.environ.pop("RATOMIZER_DOCX_EXTRA_CHANNELS", None)

    def _make_docx(self) -> Path:
        path = Path(self.tmpdir.name) / "channels.docx"
        doc = Document()
        doc.add_paragraph("Body")
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Header line"
        section.footer.paragraphs[0].text = "Footer line"
        doc.save(path)
        return path

    def test_extra_channel_blocks_are_written_when_switch_on(self):
        path = self._make_docx()
        out_dir = Path(self.tmpdir.name) / "out"
        run_atomizer_pipeline(path, out_dir)
        blocks = [
            json.loads(line)
            for line in (out_dir / "blocks.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
        channels = {b.get("content_channel") for b in blocks}
        self.assertIn("header", channels)
        self.assertIn("footer", channels)
        # 登记册应包含 header/footer 通道
        registry = json.loads((out_dir / "unextracted_registry.json").read_text(encoding="utf-8"))
        kinds = {e["kind"] for e in registry["entries"]}
        self.assertIn("header_channel", kinds)
        self.assertIn("footer_channel", kinds)

    def test_extra_channel_blocks_are_not_written_when_switch_off(self):
        os.environ["RATOMIZER_DOCX_EXTRA_CHANNELS"] = "0"
        path = self._make_docx()
        out_dir = Path(self.tmpdir.name) / "out_off"
        run_atomizer_pipeline(path, out_dir)
        blocks = [
            json.loads(line)
            for line in (out_dir / "blocks.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
        extra_channels = {b.get("content_channel") for b in blocks if b.get("content_channel")}
        self.assertEqual(extra_channels, set())
        registry = json.loads((out_dir / "unextracted_registry.json").read_text(encoding="utf-8"))
        kinds = {e["kind"] for e in registry["entries"]}
        self.assertNotIn("header_channel", kinds)
        self.assertNotIn("footer_channel", kinds)


class HtmlPipelineIntegrationTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        os.environ["RATOMIZER_ENABLE_HTML_PARSER"] = "1"

    def tearDown(self):
        self.tmpdir.cleanup()
        os.environ.pop("RATOMIZER_ENABLE_HTML_PARSER", None)

    def test_html_input_runs_when_enabled(self):
        path = Path(self.tmpdir.name) / "doc.html"
        path.write_text("<html><body><h1>Scope</h1><p>The meter shall log.</p></body></html>", encoding="utf-8")
        out_dir = Path(self.tmpdir.name) / "out"
        manifest = run_atomizer_pipeline(path, out_dir)
        self.assertEqual(manifest["input_format"], "html")
        blocks = [
            json.loads(line)
            for line in (out_dir / "blocks.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
        texts = [b["text"] for b in blocks]
        self.assertIn("Scope", texts)
        self.assertIn("The meter shall log.", texts)

    def test_html_input_disabled_by_default(self):
        os.environ["RATOMIZER_ENABLE_HTML_PARSER"] = "0"
        path = Path(self.tmpdir.name) / "doc2.html"
        path.write_text("<html><body><p>x</p></body></html>", encoding="utf-8")
        out_dir = Path(self.tmpdir.name) / "out2"
        from atomize import AtomizerInputError
        with self.assertRaises(AtomizerInputError):
            run_atomizer_pipeline(path, out_dir)


class XlsxRequirementListIntegrationTests(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.tmpdir.cleanup()
        os.environ.pop(XLSX_REQUIREMENT_LIST_SWITCH, None)

    def test_switch_on_produces_base_library_candidates(self):
        os.environ[XLSX_REQUIREMENT_LIST_SWITCH] = "1"
        path = Path(self.tmpdir.name) / "reqs.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "系统需求"
        ws.append(["子模块", "描述", "需求模版"])
        ws.append(["计量", "电表应支持事件记录", "The meter shall support event logging"])
        wb.save(path)
        wb.close()
        out_dir = Path(self.tmpdir.name) / "out"
        manifest = run_atomizer_pipeline(path, out_dir)
        self.assertIn("xlsx_requirement_list_candidates", manifest["counts"])
        self.assertTrue((out_dir / BASE_LIBRARY_CANDIDATES_FILE).exists())


if __name__ == "__main__":
    unittest.main()
