"""表格结构与单元格级需求闭环 v1（table-structure-v2）全场景测试。

覆盖规格 §九 的场景矩阵（DOCX/XLSX/PDF）：
普通 3×3、标题+表头+数据、无表头单列表、首行单格需求、数据区单格需求、
同行双需求、单格双句、合并标题、合并需求格、多级表头、X 映射矩阵、
普通 Note=mandatory、同 sheet 多表、XLSX 非 A1 起始区域、PDF cell bbox。

关键断言：
- 首行规范性文本恰好生成一个 claim；
- 同行两个需求生成两个 claim（focus 指纹互相独立）；
- 所有 source_item_id/source_cell_id 均真实存在；
- 不再生成 "1 shall support Note."；
- 每个非空 canonical cell 恰好被消费一次（accounting 硬门全零）；
- blocks.jsonl 的 block ID 序列稳定（不新增顶层 block）。
"""
from __future__ import annotations

import copy
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from ai_extract import _assert_source_references
from atomize import (
    build_atomic_candidates,
    build_table_artifacts,
    extract_docx,
    mark_doc_regions,
)
from claim_catalog import build_claim_catalog
from claim_focus import ClaimFocusError, build_claim_focus_adapter
from parsers.xlsx_parser import extract_xlsx
from requirement_kb import KnowledgeRepository
from table_structure import (
    TABLE_CELL_ITEM_SCHEMA,
    TABLE_STRUCTURE_VERSION,
    cell_context_text,
    physical_data_row_indexes,
)


KB = KnowledgeRepository.from_paths([])


def _artifacts(matrix, *, merges=None, title="T", section=("S",), **kwargs):
    return build_table_artifacts(
        matrix,
        table_id="TBL-000001",
        block_id="BLK-000002",
        order=2,
        table_title=title,
        section_path=list(section),
        knowledge_bases=KB,
        merge_ranges=merges,
        **kwargs,
    )


def _catalog(block, items, cells):
    return build_claim_catalog([block], items, table_cell_items=cells)


def _cell_audit(result):
    audit = result["meta"]["audit"]
    return {
        key: audit[key]
        for key in (
            "unconsumed_table_cell_count",
            "multi_consumed_table_cell_count",
            "dangling_table_item_reference_count",
            "dangling_table_cell_reference_count",
            "normative_context_only_count",
        )
    }


class PlainTableTests(unittest.TestCase):
    def test_plain_3x3_row_mode_closed(self) -> None:
        matrix = [
            ["Name", "Value", "Note"],
            ["Voltage", "230 V", "mandatory"],
            ["Current", "10 A", "optional"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["table_structure_version"], TABLE_STRUCTURE_VERSION)
        self.assertEqual(block["table_kind"], "parameter")
        self.assertEqual(block["leaf_mode"], "row")
        self.assertEqual(block["header_detection_status"], "inferred")
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        self.assertEqual(result["meta"]["table_structure_status"], "ok")
        self.assertTrue(all(count == 0 for count in _cell_audit(result).values()))
        claims = result["catalog"]
        self.assertEqual(len(claims), 2)
        self.assertTrue(all(claim["source_kind"] == "table_row" for claim in claims))
        # Note 列的 mandatory 不是矩阵 marker：不得生成 "X shall support Note."
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        self.assertFalse(
            any(candidate["requirement"].endswith("shall support Note.") for candidate in candidates),
            [candidate["requirement"] for candidate in candidates],
        )

    def test_note_column_mandatory_never_matrix_sentence(self) -> None:
        matrix = [
            ["No.", "Requirement", "Note"],
            ["1", "The meter shall store data.", "mandatory"],
            ["2", "The meter shall be sealed.", "mandatory"],
        ]
        block, items, cells = _artifacts(matrix)
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        sentences = [candidate["requirement"] for candidate in candidates]
        self.assertFalse(any("shall support Note" in text for text in sentences), sentences)
        self.assertFalse(any(text.startswith("1 shall") for text in sentences), sentences)
        # Note 列保持原文（row-owned，随行进 row claim 逐字文本），不产矩阵句式
        note_cells = [
            cell for cell in cells
            if cell["header_path"] == ["Note"] and cell["structural_role"] == "data"
        ]
        self.assertTrue(note_cells)
        self.assertTrue(all(cell["leaf_kind"] == "row" for cell in note_cells))

    def test_title_header_data_roles(self) -> None:
        matrix = [
            ["Electrical parameters", "Electrical parameters"],
            ["Name", "Requirement"],
            ["Voltage", "The meter shall operate at 230 V."],
        ]
        block, items, cells = _artifacts(matrix, merges=[(1, 1, 1, 2)])
        self.assertEqual(block["title_row_indexes"], [1])
        self.assertEqual(block["header_row_indexes"], [2])
        self.assertEqual(block["table_title"], "Electrical parameters")
        roles = {cell["cell_id"]: cell["structural_role"] for cell in cells}
        self.assertEqual(roles["TBL-000001-R000001-C000001"], "title")
        self.assertEqual(roles["TBL-000001-R000002-C000001"], "header")
        self.assertEqual(roles["TBL-000001-R000003-C000001"], "data")
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")


class HeaderlessTests(unittest.TestCase):
    def test_headerless_single_column_preserves_first_row(self) -> None:
        matrix = [
            ["The meter shall store daily data."],
            ["The meter shall store monthly data."],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["header_row_count"], 0)
        self.assertEqual(block["headers"], ["column_1"])
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0]["row_index"], 1)
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        self.assertEqual(len(result["catalog"]), 2)

    def test_first_row_single_normative_cell_exactly_one_claim(self) -> None:
        matrix = [
            ["The device shall comply with clause 5.", ""],
            ["Name", "Requirement"],
            ["Voltage", "230 V"],
        ]
        block, items, cells = _artifacts(matrix)
        # 首行规范性单格：不是标题（无合并证据）,留在结构里且恰好一个 claim
        self.assertEqual(block["title_row_indexes"], [])
        result = _catalog(block, items, cells)
        first_row_claims = [
            claim for claim in result["catalog"]
            if int(claim["locator"].get("row_index") or 0) == 1
        ]
        self.assertEqual(len(first_row_claims), 1)
        self.assertIn("shall comply", first_row_claims[0]["text"])
        self.assertEqual(result["meta"]["accounting_status"], "complete")

    def test_data_area_single_normative_cell_kept(self) -> None:
        matrix = [
            ["Name", "Requirement"],
            ["Voltage", "230 V"],
            ["Note", "The meter shall be protected against dust."],
        ]
        block, items, cells = _artifacts(matrix)
        plan_rows = block["leaf_plan"]["row_leaves"]
        self.assertIn(3, plan_rows)  # 单格规范性行不受"至少两个非空格"限制
        result = _catalog(block, items, cells)
        texts = [claim["text"] for claim in result["catalog"]]
        self.assertTrue(any("protected against dust" in text for text in texts))


class CellGranularityTests(unittest.TestCase):
    def test_two_obligations_same_row_two_claims(self) -> None:
        matrix = [
            ["Aspect", "Requirement A", "Requirement B"],
            [
                "Storage",
                "The meter shall store daily profiles for at least sixty days.",
                "The meter must protect stored profiles against unauthorized access.",
            ],
            [
                "Display",
                "The display shall show all segments during the diagnostic test.",
                "The display must remain readable under direct sunlight conditions.",
            ],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["table_kind"], "prose_grid")
        self.assertEqual(block["leaf_mode"], "cell")
        result = _catalog(block, items, cells)
        claims = result["catalog"]
        self.assertEqual(len(claims), 4)
        self.assertTrue(all(claim["source_kind"] == "table_cell" for claim in claims))
        # 同行两条义务 = 两个互相独立的 claim（不同 cell、不同指纹）——
        # 覆盖其中一个后另一个在账本里仍是独立 open 行
        row2_claims = [
            claim for claim in claims if int(claim["locator"].get("row_index") or 0) == 2
        ]
        self.assertEqual(len(row2_claims), 2)
        focuses = [
            build_claim_focus_adapter(claim, [block], items, cells)
            for claim in row2_claims
        ]
        self.assertNotEqual(focuses[0]["table_cell_id"], focuses[1]["table_cell_id"])
        self.assertNotEqual(
            focuses[0]["context_identity_hash"], focuses[1]["context_identity_hash"]
        )
        self.assertNotEqual(row2_claims[0]["claim_id"], row2_claims[1]["claim_id"])
        # 每个 claim 带前置标识格上下文（B1：Header=Value 形态——对象语境必须自带列头）
        for claim in row2_claims:
            context = claim["table_context"]
            self.assertEqual(context["row_header_context"], ["Aspect=Storage"])
            self.assertTrue(context["header_path"])

    def test_table_cell_focus_binds_and_cross_checks_table_title(self) -> None:
        matrix = [
            ["Interface", "GET", "SET"],
            ["Data access", "X", "X"],
        ]
        title = "Optical communication interface capabilities"
        block, items, cells = _artifacts(matrix, title=title)
        catalog = _catalog(block, items, cells)
        claim = next(
            row for row in catalog["catalog"]
            if row.get("source_kind") == "table_cell"
        )
        focus = build_claim_focus_adapter(claim, [block], items, cells)

        self.assertEqual(focus["table_title"], title)
        self.assertEqual(focus["adapter_version"], "claim-focus-adapter-v3")

        changed_claim = copy.deepcopy(claim)
        changed_claim["table_context"]["table_title"] = "Ethernet interface capabilities"
        with self.assertRaisesRegex(ClaimFocusError, "title changed"):
            build_claim_focus_adapter(changed_claim, [block], items, cells)

        changed_cells = copy.deepcopy(cells)
        for cell in changed_cells:
            cell["table_title"] = "Ethernet interface capabilities"
        with self.assertRaisesRegex(ClaimFocusError, "titles disagree"):
            build_claim_focus_adapter(claim, [block], items, changed_cells)

        changed_block = copy.deepcopy(block)
        changed_block["table_title"] = "Ethernet interface capabilities"
        with self.assertRaisesRegex(ClaimFocusError, "titles disagree"):
            build_claim_focus_adapter(claim, [changed_block], items, cells)

        other_block, other_items, other_cells = _artifacts(
            matrix, title="Ethernet interface capabilities"
        )
        other_claim = next(
            row for row in _catalog(other_block, other_items, other_cells)["catalog"]
            if row.get("source_kind") == "table_cell"
        )
        other_focus = build_claim_focus_adapter(
            other_claim, [other_block], other_items, other_cells
        )
        self.assertNotEqual(
            focus["context_identity_hash"], other_focus["context_identity_hash"]
        )

    def test_single_cell_two_sentences_two_claims(self) -> None:
        matrix = [
            ["Name", "Requirement"],
            ["General", "The meter shall store data. It must be tamper proof."],
        ]
        block, items, cells = _artifacts(matrix)
        # 多义务格：同格两条独立规范性句 → 按句切出两个 cell claim（owner=cell），
        # 不再骑墙在一个 row claim 里——覆盖其中一条后另一条在账本里仍独立 open
        self.assertEqual((block["leaf_plan"] or {}).get("multi_duty_cells"),
                         ["TBL-000001-R000002-C000002"])
        result = _catalog(block, items, cells)
        cell_claims = [
            claim for claim in result["catalog"]
            if claim["locator"].get("table_cell_id") == "TBL-000001-R000002-C000002"
        ]
        self.assertEqual(len(cell_claims), 2)
        self.assertIn("store data", cell_claims[0]["text"])
        self.assertNotIn("tamper proof", cell_claims[0]["text"])
        self.assertIn("tamper proof", cell_claims[1]["text"])
        self.assertNotIn("store data", cell_claims[1]["text"])
        self.assertEqual(result["meta"]["accounting_status"], "complete")

        # 短句同样按句归属（"It shall log. It must alarm." 不得被并成一条）
        short_matrix = [
            ["Name", "Behavior"],
            ["Logger", "It shall log. It must alarm."],
        ]
        block_s, items_s, cells_s = _artifacts(short_matrix)
        result_s = _catalog(block_s, items_s, cells_s)
        short_claims = [
            claim for claim in result_s["catalog"]
            if claim["locator"].get("table_cell_id") == "TBL-000001-R000002-C000002"
        ]
        self.assertEqual(len(short_claims), 2)
        self.assertIn("shall log", short_claims[0]["text"])
        self.assertIn("must alarm", short_claims[1]["text"])

        prose_matrix = [
            ["Topic", "Details"],
            ["The device shall retain stored data. It must survive outages.",
             "The meter shall store data securely. It must be tamper proof always."],
            ["The device shall render totals outdoors. It must stay readable.",
             "The display shall show totals clearly. It must remain readable outdoors."],
        ]
        block2, items2, cells2 = _artifacts(prose_matrix)
        # 表型判定是本用例断言前提：不是 prose_grid 必须当场失败，不得跳过断言
        self.assertEqual(block2["table_kind"], "prose_grid")
        result2 = _catalog(block2, items2, cells2)
        cell_claims = [
            claim for claim in result2["catalog"]
            if claim["locator"].get("table_cell_id") == "TBL-000001-R000002-C000002"
        ]
        self.assertEqual(len(cell_claims), 2)
        self.assertIn("store data", cell_claims[0]["text"])
        self.assertIn("tamper proof", cell_claims[1]["text"])

    def test_duplicate_sentence_in_cell_is_single_claim(self) -> None:
        # DOCX 合并把两格文本拼进同一 tc → 同句逐字节重复是拼接伪影，
        # 不是两条义务：只出一个 claim，且不算多义务格
        matrix = [
            ["Name", "Requirement", "Note"],
            ["General", "The meter shall be secure. The meter shall be secure.", "x"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual((block["leaf_plan"] or {}).get("multi_duty_cells") or [], [])
        result = _catalog(block, items, cells)
        secure_claims = [
            claim for claim in result["catalog"] if "shall be secure" in claim["text"]
        ]
        self.assertEqual(len(secure_claims), 1)
        self.assertEqual(result["meta"]["accounting_status"], "complete")

    def test_merged_normative_title_generates_claim(self) -> None:
        matrix = [
            ["The device shall comply.", "The device shall comply."],
            ["Name", "Req"],
            ["V", "shall be 230"],
        ]
        block, items, cells = _artifacts(matrix, merges=[(1, 1, 1, 2)])
        self.assertEqual(block["title_row_indexes"], [1])
        title_cells = [cell for cell in cells if cell["structural_role"] == "title"]
        self.assertEqual(len(title_cells), 1)
        self.assertEqual(title_cells[0]["column_span"], 2)
        self.assertEqual(title_cells[0]["covered_coordinates"], [[1, 2]])
        result = _catalog(block, items, cells)
        title_claims = [
            claim for claim in result["catalog"]
            if claim["locator"].get("table_cell_id") == title_cells[0]["cell_id"]
        ]
        self.assertEqual(len(title_claims), 1)
        self.assertEqual(title_claims[0]["text"], "The device shall comply.")
        self.assertEqual(result["meta"]["accounting_status"], "complete")

    def test_merged_requirement_cell_single_anchor_no_duplicates(self) -> None:
        matrix = [
            ["Name", "Req A", "Req B"],
            ["General", "The meter shall be secure.", "The meter shall be secure."],
        ]
        block, items, cells = _artifacts(matrix, merges=[(2, 2, 2, 3)])
        anchor = [cell for cell in cells if cell["cell_id"] == "TBL-000001-R000002-C000002"]
        self.assertEqual(len(anchor), 1)
        self.assertEqual(anchor[0]["column_span"], 2)
        # 覆盖坐标不得复制文本冒充多个单元格
        self.assertFalse(
            any(cell["cell_id"] == "TBL-000001-R000002-C000003" for cell in cells)
        )
        self.assertEqual(anchor[0]["covered_coordinates"], [[2, 3]])
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        claims = [
            claim for claim in result["catalog"]
            if "shall be secure" in claim["text"]
        ]
        self.assertEqual(len(claims), 1)

    def test_multi_level_headers(self) -> None:
        matrix = [
            ["Customer application process", "xDLMS Service", "xDLMS Service"],
            ["Customer application process", '"GET"', '"ACTION"'],
            ["Public customer", "X", ""],
            ["Management client", "", "X"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["header_row_count"], 2)
        self.assertEqual(block["headers"][1], 'xDLMS Service / "GET"')
        marker = next(
            cell for cell in cells
            if cell["cell_id"] == "TBL-000001-R000003-C000002"
        )
        self.assertEqual(marker["header_path"], ['xDLMS Service / "GET"'])
        # B1：前置标识格带列头（Header=Value），对象语境自证身份
        self.assertEqual(
            marker["row_header_context"],
            ["Customer application process=Public customer"],
        )


class MappingMatrixTests(unittest.TestCase):
    def test_x_matrix_cell_claims_and_atoms(self) -> None:
        matrix = [
            ["Feature", "Mode A", "Mode B", "Note"],
            ["Encryption", "X", "", "see below"],
            ["Signing", "X", "X", "free text"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["table_kind"], "mapping_matrix")
        self.assertEqual(block["leaf_mode"], "cell")
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        claims = result["catalog"]
        self.assertTrue(all(claim["source_kind"] == "table_cell" for claim in claims))
        self.assertEqual(len(claims), 3)
        by_cell = {claim["locator"]["table_cell_id"]: claim for claim in claims}
        self.assertEqual(
            by_cell["TBL-000001-R000002-C000002"]["table_context"]["row_header_context"],
            ["Feature=Encryption"],
        )
        # Note 列是 context，不成 claim
        self.assertFalse(any(cell_id.endswith("C000004") for cell_id in by_cell))
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        matrix_atoms = [
            candidate for candidate in candidates
            if candidate["requirement_type"] == "capability_matrix"
        ]
        self.assertEqual(len(matrix_atoms), 3)
        self.assertTrue(all(candidate["source_type"] == "table_cell" for candidate in matrix_atoms))
        self.assertIn(
            "Encryption shall support Mode A.",
            [candidate["requirement"] for candidate in matrix_atoms],
        )
        self.assertFalse(
            any("shall support Note" in candidate["requirement"] for candidate in candidates)
        )

    def test_price_schedule_stays_row_mode(self) -> None:
        matrix = [
            ["Lot", "Description", "Price"],
            ["1", "Single-phase meter", "120"],
            ["2", "Polyphase meter", "340"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["table_kind"], "other")
        self.assertEqual(block["leaf_mode"], "row")


class StructureGateTests(unittest.TestCase):
    """table-structure-v3 防伪句闸门：结构角色 ≠ 内容资格；矩阵事实只在
    显式二维维度证据下合成，证据不全保留原始 cell 文本（宁缺勿假）。"""

    def test_checklist_disposition_headers_never_matrix(self) -> None:
        # [Item/Status/Required] 检查表：Status/Required 是处置词表头，
        # 其列的 X 是检查结果不是能力维度——绝不合成 "Voltage shall support Status."
        matrix = [
            ["Item", "Status", "Required"],
            ["Voltage", "X", "X"],
            ["Current", "X", ""],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertNotEqual(block["table_kind"], "mapping_matrix")
        self.assertEqual(block["matrix_fact_columns"], [])
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        requirements = [candidate["requirement"] for candidate in candidates]
        self.assertFalse(any("Status" in text or "Required" in text for text in requirements))
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        # 裸词 claim 事故：marker 词本身永不成 claim
        self.assertFalse(
            any(claim["text"].strip() in ("Status", "Required") for claim in result["catalog"])
        )

    def test_sentence_subject_never_matrix_fact(self) -> None:
        # 义务句被误判为行头时，句子不是对象名——不合成 "… shall support Mode A."
        matrix = [
            ["Feature", "Mode A", "Mode B"],
            ["The device shall operate continuously at all times.", "X", ""],
        ]
        block, items, cells = _artifacts(matrix)
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        self.assertFalse(
            any(
                "shall operate" in candidate["requirement"]
                and "shall support" in candidate["requirement"]
                for candidate in candidates
            )
        )

    def test_ambiguous_header_never_synthetic_fact(self) -> None:
        # 首两行皆规范性句 → ambiguous：表头行保留结构角色（规范性句各自成 claim）
        # 但永不充当列名（column_N 合成表头不是维度证据，不得产伪句）
        matrix = [
            ["The meter shall log all events.", "The meter shall store profiles."],
            ["Voltage", "230 V"],
            ["Current", "5 A"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["header_detection_status"], "ambiguous")
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")
        texts = [claim["text"] for claim in result["catalog"]]
        self.assertTrue(any("shall log all events" in text for text in texts))
        self.assertTrue(any("shall store profiles" in text for text in texts))
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        self.assertFalse(
            any("column_" in candidate["requirement"] for candidate in candidates)
        )
        # 数据行仍闭环（合成列名下的行文本守恒）
        row_texts = [item["text"] for item in items]
        self.assertTrue(any("230 V" in text for text in row_texts))

    def test_colon_spec_header_cells_claimed(self) -> None:
        # 表头区格内 "Name: value" 规格行是义务内容，不得因角色是 header 被丢弃
        matrix = [
            ["Property", "Value"],
            ["Battery service life: 15 years", "Outputs can be assigned using a user program."],
        ]
        block, items, cells = _artifacts(matrix)
        result = _catalog(block, items, cells)
        texts = [claim["text"] for claim in result["catalog"]]
        self.assertTrue(any("15 years" in text for text in texts))
        self.assertTrue(any("user program" in text for text in texts))
        self.assertEqual(result["meta"]["accounting_status"], "complete")

    def test_rowspan_empty_cell_inherits_anchor_context(self) -> None:
        # 纵向合并：covered 空格的行头上下文继承 anchor 文本（DOCX 扁平矩阵同口径），
        # 映射矩阵不因 covered 空格降级丢事实
        matrix = [
            ["Object", "Attr1", "Attr2"],
            ["Object A", "X", ""],
            ["", "X", "X"],
        ]
        block, items, cells = _artifacts(matrix, merges=[(2, 1, 3, 1)])
        self.assertEqual(block["table_kind"], "mapping_matrix")
        result = _catalog(block, items, cells)
        r3_claims = [
            claim for claim in result["catalog"]
            if int(claim["locator"].get("row_index") or 0) == 3
        ]
        self.assertTrue(r3_claims)
        for claim in r3_claims:
            # 纵向合并继承 anchor 对象值（B1：Header=Value 形态）
            self.assertEqual(
                claim["table_context"]["row_header_context"], ["Object=Object A"]
            )

    def test_overlapping_merge_ranges_drop_geometry_keep_text(self) -> None:
        # 合并证据自相矛盾（面积相交）→ 放弃精确合并、保留全部文本、
        # 结构标 needs_review；内容守恒与结构审核分离（accounting 仍 complete）
        matrix = [
            ["H1", "H2"],
            ["A", "B"],
            ["C", "D"],
        ]
        block, items, cells = _artifacts(
            matrix, merges=[(2, 1, 3, 1), (2, 1, 3, 2)]
        )
        self.assertEqual(block["merge_evidence_status"], "dropped_conflict")
        self.assertEqual(block["merge_ranges"], [])
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        flat = block["text"] + " " + " ".join(
            " ".join(str(cell or "") for cell in row) for row in block["data_rows"]
        )
        for token in ("A", "B", "C", "D"):
            self.assertIn(token, flat)


class SourceReferenceAssertionTests(unittest.TestCase):
    def test_source_ids_must_exist(self) -> None:
        matrix = [
            ["Name", "Requirement"],
            ["Voltage", "The meter shall operate at 230 V."],
        ]
        block, items, cells = _artifacts(matrix)
        item_id = items[0]["item_id"]
        cell_id = cells[0]["cell_id"]
        _assert_source_references(
            [{"source_item_id": item_id}, {"source_cell_id": cell_id}],
            items,
            cells,
        )
        with self.assertRaises(ValueError):
            _assert_source_references(
                [{"source_item_id": "TBL-000001-R000099"}], items, cells
            )
        with self.assertRaises(ValueError):
            _assert_source_references(
                [{"source_cell_id": "TBL-000001-R000002-C000099"}], items, cells
            )

    def test_cell_context_text_never_bare(self) -> None:
        matrix = [
            ["Feature", "Mode A"],
            ["Encryption", "X"],
            ["Signing", "X"],
        ]
        block, items, cells = _artifacts(matrix)
        marker = next(cell for cell in cells if cell["leaf_kind"] == "cell")
        context_text = cell_context_text(marker)
        self.assertIn(marker["table_title"], context_text)
        self.assertIn("Encryption", context_text)
        self.assertIn("Mode A", context_text)
        self.assertTrue(context_text.endswith("= X"))


class DocxIntegrationTests(unittest.TestCase):
    def _write_docx(self, path: Path) -> None:
        document = Document()
        document.add_heading("5 Requirements", level=1)
        table = document.add_table(rows=4, cols=3)
        table.cell(0, 0).merge(table.cell(0, 2))
        table.cell(0, 0).text = "Table 1 - Electrical"
        table.cell(1, 0).text = "Name"
        table.cell(1, 1).text = "Requirement"
        table.cell(1, 2).text = "Note"
        table.cell(2, 0).text = "Voltage"
        table.cell(2, 1).text = "The meter shall operate at 230 V."
        table.cell(2, 2).text = "mandatory"
        table.cell(3, 0).text = "General"
        table.cell(3, 1).text = "The meter shall be secure."
        table.cell(3, 2).text = "The meter shall be secure."
        table.cell(3, 1).merge(table.cell(3, 2))
        document.save(path)

    def test_docx_merge_evidence_and_block_sequence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spec.docx"
            self._write_docx(path)
            blocks, items, cells = extract_docx(path)
        block_ids = [block["block_id"] for block in blocks]
        self.assertEqual(block_ids, [f"BLK-{index + 1:06d}" for index in range(len(blocks))])
        table_block = next(block for block in blocks if block["type"] == "table")
        self.assertEqual(table_block["table_structure_version"], TABLE_STRUCTURE_VERSION)
        self.assertEqual(table_block["title_row_indexes"], [1])
        self.assertTrue(table_block["merge_ranges"])
        result = build_claim_catalog(blocks, items, table_cell_items=cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        self.assertTrue(all(count == 0 for count in _cell_audit(result).values()))
        # 合并需求格：单 anchor 一个 claim，覆盖坐标不冒充
        merged_claims = [
            claim for claim in result["catalog"]
            if "shall be secure" in claim["text"]
        ]
        self.assertEqual(len(merged_claims), 1)
        # Note 列的 mandatory 不产矩阵句式
        candidates = build_atomic_candidates(blocks, items, table_cell_items=cells)
        self.assertFalse(
            any("shall support Note" in candidate["requirement"] for candidate in candidates)
        )


class XlsxRegionTests(unittest.TestCase):
    def test_multi_table_sheet_split_and_non_a1_origin(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "multi.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Specs"
            # 表 1：A1 起始（标题合并 + 表头 + 数据）
            sheet["A1"] = "Electrical"
            sheet.merge_cells("A1:B1")
            sheet.append(["Name", "Requirement"])
            sheet.append(["Voltage", "The meter shall operate at 230 V."])
            sheet.append([])
            # 表 2：非 A1 起始区域（C6 起始）
            sheet["C6"] = "Name"
            sheet["D6"] = "Requirement"
            sheet["C7"] = "Current"
            sheet["D7"] = "The meter shall measure current."
            workbook.save(path)

            blocks, items, cells = extract_xlsx(path, knowledge_bases=[], document_profile=None)

        tables = [block for block in blocks if block["type"] == "table"]
        self.assertEqual(len(tables), 2)
        first, second = tables
        self.assertEqual(first["table_title"], "Electrical")
        self.assertEqual(first["title_row_indexes"], [1])
        self.assertEqual(second["headers"], ["Name", "Requirement"])
        # 非 A1 起始：cell 的 a1_address 必须带真实 sheet 坐标
        second_cells = [
            cell for cell in cells if cell["table_block_id"] == second["block_id"]
        ]
        self.assertTrue(second_cells)
        by_text = {cell["text"]: cell for cell in second_cells}
        self.assertEqual(by_text["Name"]["a1_address"], "C6")
        self.assertEqual(by_text["Current"]["a1_address"], "C7")
        self.assertEqual(by_text["Name"]["sheet_name"], "Specs")
        result = build_claim_catalog(blocks, items, table_cell_items=cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")

    def test_excel_table_definition_region(self) -> None:
        from openpyxl.worksheet.table import Table as XlsxTable
        from openpyxl.worksheet.table import TableStyleInfo

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "tables.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Data"
            sheet.append(["Name", "Requirement"])
            sheet.append(["Voltage", "The meter shall operate at 230 V."])
            sheet.append(["Current", "The meter shall measure current."])
            xlsx_table = XlsxTable(displayName="SpecTable", ref="A1:B3")
            xlsx_table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium9")
            sheet.add_table(xlsx_table)
            workbook.save(path)

            blocks, items, cells = extract_xlsx(path, knowledge_bases=[], document_profile=None)

        tables = [block for block in blocks if block["type"] == "table"]
        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0]["header_detection_status"], "explicit")
        self.assertEqual(tables[0]["header_row_indexes"], [1])


class PdfCellGeometryTests(unittest.TestCase):
    def test_cluster_boundaries_none_means_no_evidence_only(self) -> None:
        # S13（review-2026-08-03）：链式聚类保证相邻簇中心距恒 > tolerance，
        # 旧"中心过近=证据自相矛盾→None"分支不可达；None 统一为"无证据"
        # （空输入），与调用方把 falsy 映射为 status="none" 的口径一致。
        from parsers.pdf_parser import _cluster_boundaries

        self.assertIsNone(_cluster_boundaries([]))
        # 非空输入恒返回簇中心——不存"矛盾"出口
        self.assertEqual(_cluster_boundaries([0.0, 1.0, 5.0], tolerance=2.0), [0.5, 5.0])
        self.assertEqual(_cluster_boundaries([3.0]), [3.0])
        wide = _cluster_boundaries([0.0, 2.0, 4.5], tolerance=2.0)
        self.assertIsNotNone(wide)
        # docstring 不得再把 None 描述为矛盾证据（防未来聚类变动后语义回潮）
        doc = (_cluster_boundaries.__doc__ or "")
        self.assertNotIn("自相矛盾", doc)
        self.assertIn("无证据", doc)

    def test_pdfplumber_cell_evidence_anchor_and_merge(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _FakeTable:
            # 3 列 × 2 行网格；(1,1)-(1,2) 横向合并为一个 anchor
            cells = [
                (0.0, 0.0, 100.0, 20.0),   # anchor R1C1（跨 C1-C2）
                (100.0, 0.0, 150.0, 20.0),  # R1C3
                (0.0, 20.0, 50.0, 40.0),    # R2C1
                (50.0, 20.0, 100.0, 40.0),  # R2C2
                (100.0, 20.0, 150.0, 40.0),  # R2C3
            ]

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(_FakeTable())
        self.assertEqual(status, "ok")
        self.assertIsNotNone(cell_bboxes)
        self.assertEqual(cell_bboxes[(1, 1)], [0.0, 0.0, 100.0, 20.0])
        self.assertNotIn((1, 2), cell_bboxes)
        self.assertEqual(merge_ranges, [(1, 1, 1, 2)])

    def test_pdfplumber_cell_evidence_misaligned_is_honest_none(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _BrokenTable:
            # 错位矩形：与网格边界对不齐——不得反推出跨行跨列的假 merge
            cells = [(0.0, 0.0, 50.0, 20.0), (7.3, 3.1, 60.0, 25.0)]

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(_BrokenTable())
        # M1：错位 = 证据内部矛盾（conflict），与"无几何"（none）显式区分；
        # 失败时保留文本、放弃精确几何，绝不伪造
        self.assertEqual(status, "conflict")
        self.assertIsNone(cell_bboxes)
        self.assertIsNone(merge_ranges)

    def test_pdfplumber_cell_evidence_overlap_is_honest_none(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _OverlapTable:
            cells = [
                (0.0, 0.0, 100.0, 20.0),
                (50.0, 0.0, 150.0, 20.0),  # 与上一格 x 区间重叠
                (0.0, 20.0, 50.0, 40.0),
                (50.0, 20.0, 100.0, 40.0),
                (100.0, 20.0, 150.0, 40.0),
            ]

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(_OverlapTable())
        self.assertEqual(status, "conflict")
        self.assertIsNone(cell_bboxes)
        self.assertIsNone(merge_ranges)

    def test_pdfplumber_cell_evidence_dimension_mismatch_is_none(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _FakeTable:
            cells = [
                (0.0, 0.0, 100.0, 20.0),
                (100.0, 0.0, 150.0, 20.0),
                (0.0, 20.0, 50.0, 40.0),
                (50.0, 20.0, 100.0, 40.0),
                (100.0, 20.0, 150.0, 40.0),
            ]

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(
            _FakeTable(), expected_rows=3, expected_columns=3
        )
        # 几何网格与文本矩阵维度互相矛盾 → conflict（调用方传播 needs_review）
        self.assertEqual(status, "conflict")
        self.assertIsNone(cell_bboxes)
        self.assertIsNone(merge_ranges)

    def test_pdfplumber_cell_evidence_anchor_on_covered_is_honest_none(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _AnchorOnCoveredTable:
            # 大格 R1C1 跨 C1-C2；另一个矩形把 anchor 落在其 covered 格 (1,2)——
            # anchor-on-covered 与 anchor-on-anchor 同为几何矛盾，必须整体放弃
            cells = [
                (0.0, 0.0, 100.0, 20.0),   # anchor (1,1)，covered (1,2)
                (50.0, 0.0, 100.0, 20.0),  # anchor (1,2) 落进上一格的 covered 区
                (0.0, 20.0, 50.0, 40.0),
                (50.0, 20.0, 100.0, 40.0),
            ]

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(_AnchorOnCoveredTable())
        self.assertEqual(status, "conflict")
        self.assertIsNone(cell_bboxes)
        self.assertIsNone(merge_ranges)

    def test_pdfplumber_cell_evidence_identical_bbox_deduped(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _DupTable:
            # 完全相同的 bbox 是 pdfplumber 偶发的重复输出（合法），去重后正常出区
            cells = [
                (0.0, 0.0, 50.0, 20.0),
                (0.0, 0.0, 50.0, 20.0),
                (50.0, 0.0, 100.0, 20.0),
                (0.0, 20.0, 50.0, 40.0),
                (50.0, 20.0, 100.0, 40.0),
            ]

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(_DupTable())
        self.assertEqual(status, "ok")
        self.assertIsNotNone(cell_bboxes)
        self.assertEqual(len(cell_bboxes), 4)
        self.assertIsNone(merge_ranges)

    def test_pdfplumber_cell_evidence_no_cells_is_none_not_conflict(self) -> None:
        from parsers.pdf_parser import _pdfplumber_cell_evidence

        class _EmptyTable:
            cells = []

        cell_bboxes, merge_ranges, status = _pdfplumber_cell_evidence(_EmptyTable())
        # 无几何证据 = none（如实降级），不是矛盾
        self.assertEqual(status, "none")
        self.assertIsNone(cell_bboxes)
        self.assertIsNone(merge_ranges)


class ReviewCounterexampleTests(unittest.TestCase):
    """2026-07-31 复审反例（B1/B3/B4/B5/B6/M4）——每个最小反例必须直接钉死。"""

    def test_b1_same_row_obligations_carry_object_identity(self) -> None:
        # B1：`Logger | It shall log. It must alarm.` 两条句级 claim 都必须带 Logger
        matrix = [
            ["Object", "Behavior"],
            ["Logger", "It shall log. It must alarm."],
        ]
        block, items, cells = _artifacts(matrix)
        result = _catalog(block, items, cells)
        cell_claims = [
            claim for claim in result["catalog"]
            if claim["locator"].get("table_cell_id") == "TBL-000001-R000002-C000002"
        ]
        self.assertEqual(len(cell_claims), 2)
        for claim in cell_claims:
            self.assertEqual(
                claim["table_context"]["row_header_context"], ["Object=Logger"]
            )
            self.assertIn("Logger", claim["semantic_context"])
        # 对象格被 claim 消费为身份上下文——不是"消失内容"，不计 unsignaled
        self.assertEqual(result["meta"]["audit"]["unsignaled_data_cell_count"], 0)

    def test_b3_disposition_marker_table_no_bare_x_claims(self) -> None:
        # B3：`Item | Status | Required` 的 X 是处置状态——不判矩阵、不产裸 X claim、
        # 拒收列进 needs_review，行原文逐字保留。
        # P1-1 复审：other 表无强义务信号的行不再按单元格数量登记为正式 claim
        # （"a | ok | X" 普通处置行进 claim 面正是误登记）——原文留在 context
        # 格（可定位）、计数并 needs_review
        matrix = [
            ["Item", "Status", "Required"],
            ["a", "ok", "X"],
            ["b", "fail", "X"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertNotEqual(block["table_kind"], "mapping_matrix")
        self.assertEqual(block["matrix_fact_columns"], [])
        self.assertEqual(block["matrix_rejected_marker_columns"], [2])
        result = _catalog(block, items, cells)
        self.assertFalse(
            any(
                claim["source_kind"] == "table_cell"
                and claim["eligibility"] == "claim"
                for claim in result["catalog"]
            )
        )
        self.assertFalse(
            any(
                str(claim["text"]).strip() == "X"
                and claim["eligibility"] == "claim"
                for claim in result["catalog"]
            )
        )
        # P1-1：处置行不再被登记成正式 claim
        self.assertFalse(
            any("Status" in claim["text"] and "Required" in claim["text"]
                for claim in result["catalog"])
        )
        review_candidates = [
            claim for claim in result["catalog"]
            if claim["eligibility"] == "excluded"
            and claim["exclusion"]["reason"] == "unsignaled_table_cell"
        ]
        self.assertGreater(len(review_candidates), 0)
        self.assertTrue(all(
            claim["locator"]["table_cell_id"]
            == claim["exclusion"]["evidence"]["table_cell_id"]
            for claim in review_candidates
        ))
        marker_candidates = [
            claim for claim in result["catalog"]
            if claim["eligibility"] == "excluded"
            and claim["exclusion"]["reason"] == "rejected_matrix_marker_cell"
        ]
        self.assertEqual(len(marker_candidates), 2)
        self.assertTrue(all(claim["text"] == "X" for claim in marker_candidates))
        marker_contexts = {claim["semantic_context"] for claim in marker_candidates}
        self.assertTrue(all("Required" in context for context in marker_contexts))
        self.assertTrue(any("Item=a" in context for context in marker_contexts))
        self.assertTrue(any("Item=b" in context for context in marker_contexts))
        self.assertTrue(all(
            claim["locator"]["table_cell_id"]
            == claim["exclusion"]["evidence"]["table_cell_id"]
            for claim in marker_candidates
        ))
        # 内容守恒：行格原文在 cell 面可定位（不冒充 claim），计数 + needs_review
        cell_texts = {cell["text"] for cell in cells if cell["structural_role"] == "data"}
        self.assertTrue({"a", "ok", "b", "fail", "X"} <= cell_texts)
        self.assertGreater(result["meta"]["audit"]["unsignaled_data_cell_count"], 0)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

    def test_b4_disposition_headers_no_obligation_synthesis(self) -> None:
        # B4：`Voltage shall support Requirement.` / `Voltage shall have Status set to ok.`
        # 伪句式族灭绝——处置/泛称表头不提供二维能力维度，值保留原文不合成
        from atomize import extract_valued_matrix_facts

        facts = extract_valued_matrix_facts(
            {"Object": "Voltage", "Status": "ok", "Class": "1"}
        )
        self.assertEqual([fact["predicate_header"] for fact in facts], ["Class"])

        matrix = [
            ["Object", "Requirement"],
            ["Voltage", "X"],
            ["Current", "X"],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertNotEqual(block["table_kind"], "mapping_matrix")
        candidates = build_atomic_candidates([block], items, table_cell_items=cells)
        self.assertNotIn(
            "Voltage shall support Requirement.",
            [candidate["requirement"] for candidate in candidates],
        )
        # 行原文保留（内容守恒不以合成句式为代价）——P1-1 起 other 表无强信号行
        # 不进 claim 面，经 context 格可定位 + 计数 + needs_review
        result = _catalog(block, items, cells)
        cell_texts = {cell["text"] for cell in cells if cell["structural_role"] == "data"}
        self.assertTrue({"Voltage", "Current", "X"} <= cell_texts)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

    def test_b5_sentence_shape_title_is_weak_signal_not_claim(self) -> None:
        # B5：句形说明句既不判标题、也不冒充列名、更不单独成正式 claim——
        # 弱信号进 context 计数 + ambiguous/needs_review
        matrix = [
            ["This table lists the general device capabilities.", ""],
            ["Name", "Requirement"],
            ["Logger", "The meter shall log events."],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["title_row_indexes"], [])
        self.assertEqual(block["header_detection_status"], "ambiguous")
        self.assertIn(
            "TBL-000001-R000001-C000001",
            (block["leaf_plan"] or {}).get("weak_signal_cells") or [],
        )
        result = _catalog(block, items, cells)
        self.assertFalse(
            any(
                str(claim["text"]).startswith("This table lists")
                and claim["eligibility"] == "claim"
                for claim in result["catalog"]
            )
        )
        self.assertEqual(result["meta"]["audit"]["weak_signal_context_cell_count"], 1)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

    def test_b5_pattern_sentence_without_period_is_not_column_name(self) -> None:
        # B5：无句号能力句（can be assigned）是义务句不是列名——headerless，
        # 内容成行 claim 逐字保留
        matrix = [
            ["Outputs can be assigned using a user program", "Value"],
            ["Relay output", "The meter shall expose the relay output."],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["header_row_count"], 0)
        self.assertFalse(
            any("Outputs can be assigned" in str(header) for header in block["headers"])
        )
        result = _catalog(block, items, cells)
        self.assertTrue(
            any(
                "Outputs can be assigned using a user program" in claim["text"]
                for claim in result["catalog"]
            )
        )

    def test_b5_prose_grid_explanatory_cell_stays_context(self) -> None:
        # B5：散文网格中的说明句格不单独成 cell claim（弱信号），规范性格照常闭环
        matrix = [
            ["Aspect", "Behavior", "Duty"],
            ["Logging", "The device shall log all tamper events immediately.",
             "The logger shall persist records for one year."],
            ["Alarms", "The device shall raise alarms on tamper detection.",
             "This column explains the alarm design rationale."],
        ]
        block, items, cells = _artifacts(matrix)
        self.assertEqual(block["table_kind"], "prose_grid")
        result = _catalog(block, items, cells)
        self.assertFalse(
            any(
                "design rationale" in str(claim["text"])
                and claim["eligibility"] == "claim"
                for claim in result["catalog"]
            )
        )
        self.assertTrue(
            any("persist records" in str(claim["text"]) for claim in result["catalog"])
        )
        self.assertEqual(result["meta"]["audit"]["weak_signal_context_cell_count"], 1)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

    def test_b6_conflicting_covered_text_rejects_merge_keeps_cells(self) -> None:
        # B6：被覆盖格文本与 anchor 不同 → 该 merge range 拒收、全部格保留、needs_review
        matrix = [
            ["Name", "Requirement"],
            ["Logger", "The meter shall log events."],
            ["Alarms", "The meter shall raise alarms."],
        ]
        block, items, cells = _artifacts(matrix, merges=[(2, 1, 3, 1)])
        self.assertEqual(block["merge_evidence_status"], "dropped_text_conflict")
        cell_ids = {cell["cell_id"] for cell in cells}
        self.assertIn("TBL-000001-R000003-C000001", cell_ids)
        alarms = next(cell for cell in cells if cell["cell_id"] == "TBL-000001-R000003-C000001")
        self.assertEqual(alarms["text"], "Alarms")
        result = _catalog(block, items, cells)
        self.assertTrue(
            any("Alarms" in claim["text"] for claim in result["catalog"])
        )
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

        # 正向对照：covered 格为空时合并合法保留
        ok_block, _ok_items, _ok_cells = _artifacts(
            [row[:] for row in matrix[:2]] + [["", "The meter shall raise alarms."]],
            merges=[(2, 1, 3, 1)],
        )
        self.assertEqual(ok_block["merge_evidence_status"], "available")
        self.assertEqual(ok_block["merge_ranges"], [[2, 1, 3, 1]])

    def test_m4_duplicate_sentence_keeps_span_aliases(self) -> None:
        # M4：同格重复句只出一个 claim，但被合并出现的 locator 全部保留
        matrix = [
            ["Object", "Behavior"],
            ["Logger", "It shall log events. It must alarm. It shall log events."],
        ]
        block, items, cells = _artifacts(matrix)
        result = _catalog(block, items, cells)
        cell_claims = [
            claim for claim in result["catalog"]
            if claim["locator"].get("table_cell_id") == "TBL-000001-R000002-C000002"
        ]
        self.assertEqual(len(cell_claims), 2)
        log_claim = next(claim for claim in cell_claims if "log events" in claim["text"])
        self.assertEqual(len(log_claim["span_aliases"]), 1)
        alias = log_claim["span_aliases"][0]
        text = "It shall log events. It must alarm. It shall log events."
        third_start = text.index("It shall log events.", 1)
        self.assertEqual(alias["cell_start"], third_start)
        self.assertEqual(alias["cell_end"], third_start + len("It shall log events."))
        alarm_claim = next(claim for claim in cell_claims if "alarm" in claim["text"])
        self.assertEqual(alarm_claim["span_aliases"], [])

    def test_p0_5_single_capability_cell_never_silently_closed(self) -> None:
        # P0-5 复审探针：无 caption/merge/style 证据的单格能力一律不得
        # "0 claim + 审计全零 + status=ok" 静默消失——无信号格为可定位候选
        # （context + 计数 + needs_review），colon_spec 强信号格在数据位授权成
        # claim（P1-1 表型×强信号口径）。（生产解析器对无合并表给
        # merge_ranges=[]，不是 None）
        context_probes = [
            "Configurable auxiliary output",
            "User-programmable outputs",
            "Outputs selected by the operator",
        ]
        for probe in context_probes:
            with self.subTest(probe=probe):
                matrix = [
                    ["Name", "Requirement"],
                    ["Logger", "The meter shall log events."],
                    [probe, ""],
                ]
                block, items, cells = _artifacts(matrix, merges=[])
                result = _catalog(block, items, cells)
                # 不再被旧同值启发式误判分组标题：行留在数据区，格原文可定位
                probe_cells = [cell for cell in cells if cell["text"] == probe]
                self.assertEqual(len(probe_cells), 1)
                self.assertEqual(probe_cells[0]["structural_role"], "data")
                # 审计不再是全零、状态不再是 ok
                audit = result["meta"]["audit"]
                self.assertGreater(
                    audit["unsignaled_data_cell_count"]
                    + audit["weak_signal_context_cell_count"]
                    + audit["ambiguous_structure_cell_count"],
                    0,
                )
                self.assertEqual(
                    result["meta"]["table_structure_status"], "needs_review"
                )
        # colon_spec 强信号在数据位：授权成 claim（不当候选吞掉，也不静默）
        matrix = [
            ["Name", "Requirement"],
            ["Logger", "The meter shall log events."],
            ["Battery service life: 15 years", ""],
        ]
        block, items, cells = _artifacts(matrix, merges=[])
        result = _catalog(block, items, cells)
        self.assertTrue(
            any("Battery service life: 15 years" in str(claim["text"])
                for claim in result["catalog"])
        )

    def test_p0_5_single_cell_first_row_is_ambiguous_candidate_not_title(self) -> None:
        # P0-5：首行单格（无合并证据）不静默判标题——可定位的歧义资格候选
        for probe in ("Configurable auxiliary output", "Battery service life: 15 years"):
            with self.subTest(probe=probe):
                matrix = [
                    [probe, ""],
                    ["Name", "Requirement"],
                    ["Logger", "The meter shall log events."],
                ]
                block, items, cells = _artifacts(matrix, merges=[])
                self.assertEqual(block["title_row_indexes"], [])
                self.assertEqual(block["header_detection_status"], "ambiguous")
                plan = block["leaf_plan"] or {}
                self.assertIn(
                    "TBL-000001-R000001-C000001",
                    plan.get("ambiguous_structure_cells") or [],
                )
                result = _catalog(block, items, cells)
                if probe == "Battery service life: 15 years":
                    # 类型化技术规格保留逐字 claim；结构角色仍待审。
                    self.assertTrue(
                        any(probe in str(claim["text"]) for claim in result["catalog"])
                    )
                else:
                    self.assertFalse(
                        any(
                            probe in str(claim["text"])
                            and claim["eligibility"] == "claim"
                            for claim in result["catalog"]
                        )
                    )
                    candidates = [
                        claim for claim in result["catalog"]
                        if probe in str(claim["text"])
                    ]
                    self.assertEqual(len(candidates), 1)
                    self.assertEqual(candidates[0]["eligibility"], "excluded")
                    self.assertEqual(
                        candidates[0]["exclusion"]["reason"],
                        "ambiguous_table_structure",
                    )
                    evidence = candidates[0]["exclusion"]["evidence"]
                    self.assertEqual(
                        evidence["table_cell_id"],
                        candidates[0]["locator"]["table_cell_id"],
                    )
                    self.assertEqual(
                        evidence["table_block_id"],
                        candidates[0]["locator"]["block_id"],
                    )
                self.assertGreaterEqual(
                    result["meta"]["audit"]["ambiguous_structure_cell_count"], 1
                )
                self.assertEqual(
                    result["meta"]["table_structure_status"], "needs_review"
                )

    def test_p0_5_one_by_one_table_not_silent(self) -> None:
        # P0-5：1×1 退化表同样不得静默关闭为表头
        block, items, cells = _artifacts([["Configurable auxiliary output"]], merges=[])
        self.assertEqual(block["header_detection_status"], "ambiguous")
        result = _catalog(block, items, cells)
        self.assertGreaterEqual(
            result["meta"]["audit"]["ambiguous_structure_cell_count"], 1
        )
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

    def test_p0_5_missing_merge_evidence_never_grants_group_header(self) -> None:
        # v5：[]（确认无合并）与 None（当前解析无几何）都没有正向全宽 merge 证据；
        # 二者均不得启用旧同值启发式。旧 artifact 由版本门迁移，不在解析层兼容。
        matrix = [
            ["Name", "Requirement"],
            ["Logger", "The meter shall log events."],
            ["User-programmable outputs", ""],
        ]
        modern, _m_items, _m_cells = _artifacts(matrix, merges=[])
        unavailable, _l_items, unavailable_cells = _artifacts(
            matrix, merges=None, source_format="pdf"
        )
        modern_plan = modern["leaf_plan"] or {}
        unavailable_plan = unavailable["leaf_plan"] or {}
        self.assertIn(
            "TBL-000001-R000003-C000001",
            modern_plan.get("unsignaled_data_cells") or [],
        )
        self.assertIn(
            "TBL-000001-R000003-C000001",
            unavailable_plan.get("unsignaled_data_cells") or [],
        )
        self.assertEqual(unavailable["merge_evidence_status"], "unavailable")
        result = _catalog(unavailable, _l_items, unavailable_cells)
        self.assertEqual(result["meta"]["table_structure_status"], "needs_review")

    def test_p0_5_modal_first_row_single_cell_still_claims(self) -> None:
        # P0-5 对照：首行单格 modal 义务句不回收——规范性内容绝不静默丢失
        matrix = [
            ["The device shall comply with clause 5.", ""],
            ["Name", "Requirement"],
            ["Voltage", "230 V"],
        ]
        block, items, cells = _artifacts(matrix, merges=[])
        result = _catalog(block, items, cells)
        first_row_claims = [
            claim for claim in result["catalog"]
            if int(claim["locator"].get("row_index") or 0) == 1
        ]
        self.assertEqual(len(first_row_claims), 1)
        self.assertIn("shall comply", first_row_claims[0]["text"])

    def test_p0_4_participle_headers_never_dimension(self) -> None:
        # P0-4 复审复现：`Feature | Supported` → "Encryption shall support Supported."、
        # `Item | Approved` → "Design shall support Approved." 均为确定性伪造——
        # 正向维度证据下分词式普通词表头不成维度，未知二维表保留原文 + needs_review
        for headers, rows, subject in (
            (["Feature", "Supported"], [["Encryption", "X"], ["Signing", "X"]], "Encryption"),
            (["Item", "Approved"], [["Design", "X"], ["Release", "X"]], "Design"),
            (["Feature", "Supported by ISO 9001"], [["Encryption", "X"], ["Signing", "X"]], "Encryption"),
            (["Feature", "Approved 2024"], [["Encryption", "X"], ["Signing", "X"]], "Encryption"),
            (["Feature", "Description v2"], [["Encryption", "X"], ["Signing", "X"]], "Encryption"),
            (["Feature", "Supported Mode"], [["Encryption", "X"], ["Signing", "X"]], "Encryption"),
        ):
            with self.subTest(headers=headers):
                matrix = [headers, *rows]
                block, items, cells = _artifacts(matrix, merges=[])
                self.assertNotEqual(block["table_kind"], "mapping_matrix")
                self.assertEqual(block["matrix_fact_columns"], [])
                self.assertEqual(block["matrix_dimension_evidence"], {})
                candidates = build_atomic_candidates([block], items, table_cell_items=cells)
                fabricated = f"{subject} shall support {headers[1]}."
                self.assertNotIn(
                    fabricated,
                    [candidate["requirement"] for candidate in candidates],
                )
                # 原文保留（格可定位）+ 拒收列/未消费计数 → needs_review
                cell_texts = {cell["text"] for cell in cells}
                self.assertIn(subject, cell_texts)
                result = _catalog(block, items, cells)
                self.assertEqual(
                    result["meta"]["table_structure_status"], "needs_review"
                )

    def test_p0_4_positive_dimension_shapes_accepted(self) -> None:
        # v5 正向面：只有受控操作、明确轴成员或限定操作路径可驱动合成。
        from table_structure import matrix_dimension_tag

        accepted = {
            '"GET"': "operation",
            "xDLMS Service": "qualified_operation",
            'xDLMS Service / "GET"': "qualified_operation",
            "Mode A": "axis_member",
            "Attr1": "axis_member",
            "xDLMS Service: GET": "qualified_operation",
        }
        for header, tag in accepted.items():
            with self.subTest(header=header):
                self.assertEqual(matrix_dimension_tag(header), tag)
        for rejected in (
            "Supported", "Approved", "Enabled", "Status", "Required", "Note",
            "column_2", "X", "2024", "Mechanism", "Supported by ISO 9001",
            "Approved 2024", "Description v2", "Supported Mode", "",
        ):
            with self.subTest(header=rejected):
                self.assertIsNone(matrix_dimension_tag(rejected))

    def test_colon_spec_uses_typed_eligibility_not_generic_shape(self) -> None:
        metadata = ("Owner: QA", "Reference: ISO", "Revision: 2024", "Status: OK")
        for value in metadata:
            with self.subTest(value=value):
                block, items, cells = _artifacts(
                    [["Heading", "Text"], ["Info", value]], merges=[]
                )
                result = _catalog(block, items, cells)
                self.assertFalse(
                    any(value in str(claim["text"]) for claim in result["catalog"])
                )
                self.assertEqual(result["meta"]["table_structure_status"], "ok")

        for value in ("Battery service life: 15 years", "Protection degree: IP54"):
            with self.subTest(value=value):
                block, items, cells = _artifacts(
                    [["Heading", "Text"], ["Info", value]], merges=[]
                )
                result = _catalog(block, items, cells)
                self.assertTrue(
                    any(value in str(claim["text"]) for claim in result["catalog"])
                )

        for value in (
            "Mounting: DIN",
            "Memory size: 16 MB",
            "Communication profile: ABC",
        ):
            with self.subTest(value=value):
                block, items, cells = _artifacts(
                    [["Heading", "Text"], ["Info", value]], merges=[]
                )
                cell_id = next(
                    cell["cell_id"] for cell in cells if cell["text"] == value
                )
                self.assertIn(
                    cell_id,
                    block["leaf_plan"]["untyped_colon_spec_cells"],
                )
                result = _catalog(block, items, cells)
                candidates = [
                    claim for claim in result["catalog"]
                    if claim["text"] == value
                ]
                self.assertEqual(len(candidates), 1)
                self.assertEqual(candidates[0]["eligibility"], "excluded")
                self.assertEqual(
                    candidates[0]["exclusion"]["reason"],
                    "untyped_colon_spec_cell",
                )
                self.assertIn(value, candidates[0]["semantic_context"])
                self.assertEqual(
                    candidates[0]["locator"]["table_cell_id"], cell_id
                )
                self.assertGreaterEqual(
                    result["meta"]["audit"]["untyped_colon_spec_cell_count"],
                    1,
                )
                self.assertEqual(
                    result["meta"]["table_structure_status"], "needs_review"
                )

    def test_merged_technical_spec_title_is_not_silently_context_only(self) -> None:
        value = "Battery service life: 15 years"
        block, items, cells = _artifacts(
            [[value, ""], ["Name", "Value"], ["A", "B"]],
            merges=[(1, 1, 1, 2)],
        )
        result = _catalog(block, items, cells)
        self.assertTrue(
            any(value in str(claim["text"]) for claim in result["catalog"])
        )


class MixedFactColumnConservationTests(unittest.TestCase):
    """I4 复审反例（2026-08-03 清单）：mixed 表的事实列必须按实际 (row,column)
    cell-leaf 坐标剔除——非 marker 文本格（"optional"）保留在 row claim，
    有且仅有一个 owner；消费审计不得靠坐标假通过。"""

    _MATRIX = [
        ["Attribute", "Value", "Read", "Write"],
        ["Voltage", "230 V", "required", "optional"],
        ["Current", "5 A", "mandatory", "x"],
    ]

    def test_non_marker_fact_cell_owned_by_row_claim_exactly_once(self) -> None:
        block, items, cells = _artifacts(self._MATRIX)
        self.assertEqual(block["table_kind"], "parameter")
        self.assertEqual(block["leaf_mode"], "mixed")
        # Read/Write 是受控操作轴名的正向维度证据列（0-based 2/3）
        self.assertEqual(block["matrix_fact_columns"], [2, 3])
        result = _catalog(block, items, cells)
        self.assertEqual(result["meta"]["accounting_status"], "complete")
        self.assertTrue(all(count == 0 for count in _cell_audit(result).values()))
        row_claims = [
            claim for claim in result["catalog"] if claim["source_kind"] == "table_row"
        ]
        cell_claims = [
            claim for claim in result["catalog"] if claim["source_kind"] == "table_cell"
        ]
        self.assertEqual(len(row_claims), 2)
        # marker 格 required/mandatory/x 各自由 cell claim 闭环
        self.assertEqual(len(cell_claims), 3)
        voltage_row = next(
            claim for claim in row_claims if claim["locator"]["row_index"] == 2
        )
        # marker 格按坐标剔除出行文本；非 marker 的 optional 不得株连剔除——
        # 它由 row claim 承载（唯一的 owner），整列剔除时代码会把它静默丢掉
        self.assertNotIn("required", voltage_row["text"])
        self.assertIn("Write=optional", voltage_row["text"])
        # optional 非 marker：不成 cell claim，也不被任何第二条 claim 重复承载
        self.assertFalse(any("optional" in claim["text"] for claim in cell_claims))
        current_row = next(
            claim for claim in row_claims if claim["locator"]["row_index"] == 3
        )
        self.assertNotIn("mandatory", current_row["text"])
        self.assertNotIn("Write=x", current_row["text"])

    def test_audit_fails_closed_when_excluded_cell_loses_owner(self) -> None:
        # 审计不得靠坐标假通过：leaf plan 把格排除出行文本（multi_duty 逐格排除），
        # 但该格没有任何 owner（非 cell leaf/候选/context）→ unconsumed 硬失败。
        # 旧口径只查"行在 row_leaf_indexes"即计消费——本用例在修复前应假通过。
        block, items, cells = _artifacts(self._MATRIX)
        block["leaf_plan"]["multi_duty_cells"] = ["TBL-000001-R000002-C000004"]
        result = _catalog(block, items, cells)
        audit = _cell_audit(result)
        self.assertGreater(audit["unconsumed_table_cell_count"], 0)
        self.assertNotEqual(result["meta"]["accounting_status"], "complete")


class PhysicalDataRowIndexTests(unittest.TestCase):
    """S12：表格块第 N 个数据行 → 物理行号（1-based）的统一推导。"""

    def test_structure_indexes_derivation(self) -> None:
        block = {
            "rows": 5,
            "title_row_indexes": [1],
            "header_row_indexes": [2],
            "data_rows": [["a"], ["b"], ["c"]],
        }
        self.assertEqual(physical_data_row_indexes(block), [3, 4, 5])

    def test_non_contiguous_title_header_indexes(self) -> None:
        # 标题/表头不连续（标题 1、表头 2 与 4）时连续偏移公式必然错位
        block = {
            "rows": 6,
            "title_row_indexes": [1],
            "header_row_indexes": [2, 4],
            "data_rows": [["a"], ["b"], ["c"]],
        }
        self.assertEqual(physical_data_row_indexes(block), [3, 5, 6])

    def test_legacy_block_falls_back_to_contiguous_formula(self) -> None:
        # 旧产物无结构索引：退回 header_row_count + 数据区偏移的历史口径
        block = {
            "headers": ["H1", "H2"],
            "header_row_count": 1,
            "data_rows": [["a", "b"], ["c", "d"]],
        }
        self.assertEqual(physical_data_row_indexes(block), [2, 3])

    def test_legacy_block_with_title_indexes_uses_legacy_formula(self) -> None:
        block = {
            "headers": ["H1"],
            "header_row_count": 1,
            "title_row_indexes": [1],
            "data_rows": [["a"], ["b"]],
        }
        self.assertEqual(physical_data_row_indexes(block), [3, 4])


class BlockSequenceStabilityTests(unittest.TestCase):
    def test_block_id_sequence_stable_with_cells(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "seq.docx"
            document = Document()
            document.add_heading("1 Scope", level=1)
            document.add_paragraph("The meter shall be tested.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Requirement"
            table.cell(1, 0).text = "Voltage"
            table.cell(1, 1).text = "The meter shall operate at 230 V."
            document.add_paragraph("Trailing clause.")
            document.save(path)
            blocks, items, cells = extract_docx(path)

        self.assertEqual(
            [block["block_id"] for block in blocks],
            ["BLK-000001", "BLK-000002", "BLK-000003", "BLK-000004"],
        )
        types = [block["type"] for block in blocks]
        self.assertEqual(types, ["heading", "paragraph", "table", "paragraph"])
        # 行/格不升格为顶层 block
        self.assertTrue(all(item["table_block_id"] == "BLK-000003" for item in items))
        self.assertTrue(all(cell["table_block_id"] == "BLK-000003" for cell in cells))
        self.assertTrue(all(cell["schema"] == TABLE_CELL_ITEM_SCHEMA for cell in cells))


if __name__ == "__main__":
    unittest.main()
