from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def write_synthetic_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("ABNT 2022 all rights reserved")
    document.add_heading("Scope", level=1)
    document.add_paragraph("The meter shall support xDLMS GET service.")
    document.add_paragraph("5.1 Security requirements")
    list_item = document.add_paragraph("Display all segments", style="List Paragraph")
    add_numbering(list_item)

    document.add_paragraph("Table 1 - Services xDLMS")
    matrix = document.add_table(rows=3, cols=3)
    matrix.cell(0, 0).text = "Customer application process"
    matrix.cell(0, 1).text = "xDLMS Service"
    matrix.cell(0, 2).text = "xDLMS Service"
    matrix.cell(1, 0).text = "Customer application process"
    matrix.cell(1, 1).text = "GET"
    matrix.cell(1, 2).text = "ACTION"
    matrix.cell(2, 0).text = "Public customer"
    matrix.cell(2, 1).text = "X"
    matrix.cell(2, 2).text = ""

    document.add_paragraph("Table 2 - COSEM Clock object")
    cosem = document.add_table(rows=3, cols=6)
    headers = ["#", "Object/attribute name", "CL", "Type", "Value", "Access rights RC/PC/SC/LC"]
    for index, header in enumerate(headers):
        cosem.cell(0, index).text = header
    for index, value in enumerate(["1", "Clock", "8", "", "0-0:1.0.0.255", ""]):
        cosem.cell(1, index).text = value
    for index, value in enumerate(["2", "logical_name", "", "octet-string[6]", "0000010000FF", "R-/--/R-/RW"]):
        cosem.cell(2, index).text = value

    document.add_paragraph("Table 3 - Events")
    events = document.add_table(rows=2, cols=4)
    event_headers = ["Group number", "Subgroup number", "Event number", "Description of the event"]
    for index, header in enumerate(event_headers):
        events.cell(0, index).text = header
    for index, value in enumerate(["1", "10", "1", "Power down"]):
        events.cell(1, index).text = value

    document.save(path)


def write_minimal_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Scope", level=1)
    document.add_paragraph("The meter shall support xDLMS GET service.")
    document.add_paragraph("Table 1 - xDLMS services")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Customer application process"
    table.cell(0, 1).text = "xDLMS Service: GET"
    table.cell(1, 0).text = "Public customer"
    table.cell(1, 1).text = "X"
    document.save(path)


def add_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
