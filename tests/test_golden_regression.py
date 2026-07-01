from __future__ import annotations

import json
import tempfile
import unittest
from collections import Counter
from pathlib import Path

import atomize
from docx import Document
from io_utils import read_jsonl


ROOT = Path(__file__).resolve().parents[1]
GOLDEN = ROOT / "golden_sets" / "abnt_nbr_16968_v5" / "golden_summary.json"
CURRENT_OUTPUT = ROOT / "out" / "abnt_nbr_16968_atomizer_v5"


class GoldenRegressionTests(unittest.TestCase):
    def setUp(self) -> None:
        if not CURRENT_OUTPUT.exists():
            self.skipTest(f"Golden output directory not found: {CURRENT_OUTPUT}")
        self.golden = json.loads(GOLDEN.read_text(encoding="utf-8"))
        self.manifest = json.loads((CURRENT_OUTPUT / "manifest.json").read_text(encoding="utf-8"))
        self.quality = json.loads((CURRENT_OUTPUT / "quality_report.json").read_text(encoding="utf-8"))
        self.atomic = read_jsonl(CURRENT_OUTPUT / "atomic_requirements.jsonl")

    def test_manifest_counts_match_golden_baseline(self) -> None:
        for key, expected in self.golden["counts"].items():
            if key == "cosem_object_instances":
                continue
            self.assertEqual(self.manifest["counts"].get(key), expected, key)

    def test_requirement_type_distribution_matches_golden_baseline(self) -> None:
        actual = Counter(row["requirement_type"] for row in self.atomic)
        self.assertEqual(dict(actual), self.golden["requirement_type_counts"])

    def test_source_type_distribution_matches_golden_baseline(self) -> None:
        actual = Counter(row["source_type"] for row in self.atomic)
        self.assertEqual(dict(actual), self.golden["source_type_counts"])

    def test_quality_coverage_matches_golden_baseline(self) -> None:
        self.assertEqual(self.quality["coverage"], self.golden["coverage"])

    def test_representative_requirements_are_present(self) -> None:
        for expected in self.golden["representative_requirements"]:
            matches = [
                row
                for row in self.atomic
                if row.get("requirement_type") == expected["requirement_type"]
                and row.get("object") == expected["object"]
                and expected["requirement_contains"] in row.get("requirement", "")
            ]
            self.assertTrue(matches, expected)


class FreshPipelineRegressionTests(unittest.TestCase):
    def test_minimal_docx_pipeline_generates_stable_outputs(self) -> None:
        self.assertTrue(hasattr(atomize, "run_atomizer_pipeline"))

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "minimal_standard.docx"
            out_dir = tmp_path / "out"
            self.write_minimal_docx(input_path)

            manifest = atomize.run_atomizer_pipeline(input_path, out_dir, chunk_chars=800)

            atomic = read_jsonl(out_dir / "atomic_requirements.jsonl")
            quality = json.loads((out_dir / "quality_report.json").read_text(encoding="utf-8"))

        self.assertEqual(manifest["counts"]["blocks"], 4)
        self.assertEqual(manifest["counts"]["table_items"], 1)
        self.assertEqual(manifest["counts"]["atomic_requirements"], 2)
        self.assertEqual(manifest["counts"]["llm_tasks"], 1)
        self.assertEqual(
            [row["requirement_type"] for row in atomic],
            ["communication", "capability_matrix"],
        )
        self.assertEqual(
            [row["requirement"] for row in atomic],
            [
                "The meter shall support xDLMS GET service.",
                "Public customer shall support xDLMS Service: GET.",
            ],
        )
        self.assertEqual(atomic[0]["req_id"], "AREQ-000001")
        self.assertEqual(atomic[1]["req_id"], "AREQ-000002")
        self.assertRegex(atomic[0]["stable_req_id"], r"^SREQ-[0-9A-F]{16}$")
        self.assertRegex(atomic[1]["stable_req_id"], r"^SREQ-[0-9A-F]{16}$")
        self.assertEqual(quality["coverage"]["body_table_candidate_ratio"], 1.0)
        self.assertEqual(quality["coverage"]["domain_table_candidate_ratio"], 0.0)

    def write_minimal_docx(self, path: Path) -> None:
        document = Document()
        document.add_heading("Scope", level=1)
        document.add_paragraph("The meter shall support xDLMS GET service.")
        document.add_paragraph("Table 1 - xDLMS services")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Customer application process"
        table.cell(0, 1).text = "xDLMS Service: GET"
        table.cell(1, 0).text = "Public customer"
        table.cell(1, 1).text = "X"
        document.save(path)


if __name__ == "__main__":
    unittest.main()
