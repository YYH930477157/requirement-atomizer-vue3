from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_table_parser import parse_docx_table
from atomize import extract_docx
from table_dispositions import build_table_cell_dispositions


def _set_row_grid_offset(row, *, before: int = 0, after: int = 0) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if before:
        node = OxmlElement("w:gridBefore")
        node.set(qn("w:val"), str(before))
        tr_pr.append(node)
    if after:
        node = OxmlElement("w:gridAfter")
        node.set(qn("w:val"), str(after))
        tr_pr.append(node)


def _set_vmerge_continue(cell) -> None:
    # A w:vMerge without w:val defaults to "continue" in OOXML; row 1 has no
    # restart anchor, so the continuation has no canonical anchor (merge_conflict).
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:vMerge")
    tc_pr.append(node)


class DocxTablePhysicalParserTests(unittest.TestCase):
    def test_grid_before_and_after_preserve_physical_columns(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = "H1"
        table.rows[0].cells[1].text = "H2"
        table.rows[0].cells[2].text = "H3"
        table.rows[0].cells[3].text = "H4"

        # The second physical row starts after one omitted grid column and ends
        # before one omitted grid column. Remove the unused edge cells so the
        # OOXML shape matches the declared offsets.
        second = table.rows[1]
        second.cells[1].text = "A"
        second.cells[2].text = "B"
        second._tr.remove(second.cells[3]._tc)
        second._tr.remove(second.cells[0]._tc)
        _set_row_grid_offset(second, before=1, after=1)

        parsed = parse_docx_table(table)

        self.assertFalse(parsed.parse_incomplete)
        self.assertEqual(parsed.width, 4)
        self.assertEqual(parsed.matrix[1], ["", "A", "B", ""])
        self.assertEqual(parsed.cells[(2, 2)].text, "A")
        self.assertEqual(parsed.cells[(2, 3)].text, "B")

    def test_rectangular_merge_has_one_anchor_and_all_covered_coordinates(self) -> None:
        document = Document()
        table = document.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Header"
        table.cell(0, 1).text = "Value"
        table.cell(0, 2).text = "Note"
        merged = table.cell(1, 1).merge(table.cell(2, 2))
        merged.text = "The meter shall retain records."

        parsed = parse_docx_table(table)

        self.assertEqual(parsed.merge_ranges, [(2, 2, 3, 3)])
        anchor = parsed.cells[(2, 2)]
        self.assertEqual(anchor.text, "The meter shall retain records.")
        self.assertEqual(
            anchor.covered_coordinates,
            ((2, 3), (3, 2), (3, 3)),
        )
        self.assertNotIn((2, 3), parsed.cells)
        self.assertNotIn((3, 2), parsed.cells)
        self.assertNotIn((3, 3), parsed.cells)

    def test_nested_table_is_independent_and_outer_cell_keeps_only_own_text(self) -> None:
        document = Document()
        outer = document.add_table(rows=1, cols=1)
        cell = outer.cell(0, 0)
        cell.paragraphs[0].text = "Outer requirement context"
        nested = cell.add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "Parameter"
        nested.cell(0, 1).text = "Value"
        nested.cell(1, 0).text = "Voltage"
        nested.cell(1, 1).text = "230 V"

        parsed = parse_docx_table(outer)

        self.assertEqual(parsed.cells[(1, 1)].text, "Outer requirement context")
        self.assertEqual(len(parsed.nested_tables), 1)
        nested_ref = parsed.nested_tables[0]
        self.assertEqual(nested_ref.parent_coordinate, (1, 1))
        self.assertEqual(nested_ref.table.matrix[1], ["Voltage", "230 V"])

    def test_paragraph_list_and_manual_break_boundaries_are_preserved(self) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        first = cell.paragraphs[0]
        first.text = "The device shall support:"
        bullet = cell.add_paragraph(style="List Bullet")
        bullet.add_run("remote reading")
        bullet.add_run().add_break(WD_BREAK.LINE)
        bullet.add_run("local reading")

        parsed = parse_docx_table(table)
        content = parsed.cells[(1, 1)].content

        self.assertEqual(len(content.paragraphs), 2)
        self.assertEqual(content.paragraphs[1].list_level, 0)
        self.assertEqual(content.paragraphs[1].text, "remote reading\nlocal reading")
        self.assertEqual(
            parsed.cells[(1, 1)].text,
            "The device shall support:\nremote reading\nlocal reading",
        )

    def test_style_evidence_reads_the_cells_own_paragraph_objects(self) -> None:
        # FIX 2 回归护栏：_parse_cell_content 返回的 paragraph_objects 必须继续
        # 喂给 _cell_style_evidence（消除调用方的二次 Paragraph 重建），样式证据
        # 逐字节不变——粗体/对齐仍来自单元格真实段落。
        document = Document()
        table = document.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        run = paragraph.add_run("Bold header")
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parsed = parse_docx_table(table)

        evidence = parsed.cells[(1, 1)].style_evidence
        self.assertTrue(evidence["bold"])
        self.assertEqual(len(evidence["paragraph_alignments"]), 1)
        self.assertIn("CENTER", evidence["paragraph_alignments"][0])

    def test_reconciled_row_width_conflict_is_audit_note_not_blocking(self) -> None:
        # 声明 tblGrid(2 列)窄于实际行宽(4 列)：解析器确定性调和——宽取
        # max(declared, observed)、矩阵按调和宽补齐、无内容丢失。冲突保留为
        # 审计注记（reconciled=True），但不再置 parse_incomplete，整表不再
        # 因此被强制 review（review 面积随表面积而非真实歧义增长）。
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        _set_row_grid_offset(table.rows[1], before=2)

        parsed = parse_docx_table(table)

        self.assertFalse(parsed.parse_incomplete)
        self.assertEqual(parsed.width, 4)
        self.assertEqual(parsed.parse_incomplete_reason["code"], "row_width_conflict")
        issues = parsed.parse_incomplete_reason["issues"]
        self.assertEqual(len(issues), 1)
        self.assertEqual(issues[0]["code"], "row_width_conflict")
        self.assertTrue(issues[0].get("reconciled"))
        self.assertEqual(parsed.parse_incomplete_reason.get("reconciled"), True)
        self.assertIn("C", parsed.raw_text)
        self.assertIn("D", parsed.raw_text)

    def test_row_narrower_than_declared_grid_reconciles_without_blocking(self) -> None:
        # 镜像情形：行窄于声明 tblGrid——调和走补齐路径（宽度保持声明值、行尾
        # 补空列），同样零丢失，不阻塞；冲突仍作为审计注记保留。
        document = Document()
        table = document.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 2).text = "C"
        second = table.rows[1]
        second.cells[0].text = "Only"
        second._tr.remove(second.cells[2]._tc)
        second._tr.remove(second.cells[1]._tc)

        parsed = parse_docx_table(table)

        self.assertFalse(parsed.parse_incomplete)
        self.assertEqual(parsed.width, 3)
        self.assertEqual(parsed.parse_incomplete_reason["code"], "row_width_conflict")
        self.assertEqual(parsed.matrix[1], ["Only", "", ""])
        self.assertIn("Only", parsed.raw_text)

    def test_merge_conflict_without_anchor_still_blocks_parse(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        _set_vmerge_continue(table.cell(1, 0))

        parsed = parse_docx_table(table)

        self.assertTrue(parsed.parse_incomplete)
        self.assertEqual(parsed.parse_incomplete_reason["code"], "merge_conflict")
        self.assertFalse(
            parsed.parse_incomplete_reason["issues"][0].get("reconciled")
        )

    def test_reconciled_width_conflict_table_does_not_force_review_dispositions(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Parameter"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Voltage"
        table.cell(1, 1).text = "230 V"
        _set_row_grid_offset(table.rows[1], before=2)

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "width-conflict.docx"
            document.save(path)
            blocks, _items, cells = extract_docx(path)

        self.assertFalse(blocks[0]["parse_incomplete"])
        self.assertEqual(
            blocks[0]["parse_incomplete_reason"]["code"], "row_width_conflict"
        )
        dispositions = build_table_cell_dispositions(blocks, cells)
        self.assertTrue(dispositions)
        self.assertFalse(any(
            "parse_incomplete" in str(entry)
            for row in dispositions
            for entry in row["evidence"]
        ))

    def test_merge_conflict_table_routes_every_cell_to_review(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Parameter"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Voltage"
        table.cell(1, 1).text = "230 V"
        _set_vmerge_continue(table.cell(1, 0))

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "merge-conflict.docx"
            document.save(path)
            blocks, _items, cells = extract_docx(path)

        self.assertTrue(blocks[0]["parse_incomplete"])
        self.assertEqual(blocks[0]["parse_incomplete_reason"]["code"], "merge_conflict")
        dispositions = build_table_cell_dispositions(blocks, cells)
        self.assertTrue(dispositions)
        self.assertTrue(all(row["disposition"] == "review" for row in dispositions))
        self.assertTrue(all(
            any(str(entry).startswith("parse_incomplete:") for entry in row["evidence"])
            for row in dispositions
        ))

    def test_extract_docx_publishes_physical_content_and_nested_cell_sources(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Parameter"
        table.cell(0, 1).text = "Requirement"
        table.cell(1, 0).text = "Voltage"
        requirement = table.cell(1, 1)
        requirement.paragraphs[0].text = "The meter shall operate at 230 V."
        nested = requirement.add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "Mode"
        nested.cell(0, 1).text = "Value"
        nested.cell(1, 0).text = "Nominal"
        nested.cell(1, 1).text = "230 V"

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.docx"
            document.save(path)
            blocks, items, cells = extract_docx(path)

        self.assertEqual([block["block_id"] for block in blocks], ["BLK-000001"])
        self.assertEqual(blocks[0]["docx_table_physical_version"], "docx-table-physical-v2")
        self.assertEqual(len(blocks[0]["nested_tables"]), 1)
        self.assertEqual(blocks[0]["nested_tables"][0]["table_id"], "TBL-000001-N001")
        outer = next(cell for cell in cells if cell["cell_id"].endswith("R000002-C000002"))
        self.assertEqual(
            outer["content_paragraphs"][0]["text"],
            "The meter shall operate at 230 V.",
        )
        self.assertEqual(outer["nested_table_ids"], ["TBL-000001-N001"])
        nested_cells = [cell for cell in cells if cell["table_id"] == "TBL-000001-N001"]
        self.assertTrue(any(cell["text"] == "230 V" for cell in nested_cells))
        self.assertTrue(any(item["table_id"] == "TBL-000001-N001" for item in items))


    def test_nested_tables_in_two_cells_get_distinct_table_global_ordinals(self) -> None:
        # 2026-08-05 Kimi 高危 #2：两个单元格各含一个嵌套表时，原按单元格
        # enumerate(start=1) 都得 ordinal=1 → 嵌套表 ID 与其 cell ID 碰撞 →
        # conservation 审计 hard-fail 致整次 atomize 失败。序号须表级全局唯一。
        document = Document()
        outer = document.add_table(rows=1, cols=2)
        for col in range(2):
            cell = outer.cell(0, col)
            nested = cell.add_table(rows=2, cols=2)
            nested.cell(0, 0).text = "Parameter"
            nested.cell(0, 1).text = "Value"
            nested.cell(1, 0).text = f"Field{col + 1}"
            nested.cell(1, 1).text = f"{col + 1}"
        parsed = parse_docx_table(outer)
        self.assertEqual(len(parsed.nested_tables), 2)
        self.assertEqual(
            sorted(ref.ordinal for ref in parsed.nested_tables),
            [1, 2],
        )
        self.assertEqual(
            sorted(ref.parent_coordinate for ref in parsed.nested_tables),
            [(1, 1), (1, 2)],
        )

    def test_extract_docx_two_nested_tables_in_distinct_cells_do_not_collide(self) -> None:
        # 端到端：两单元格各含嵌套表 → N001/N002 两个独立嵌套表，cell_id 全唯一、
        # atomize 不因 ID 碰撞而失败。
        document = Document()
        outer = document.add_table(rows=1, cols=2)
        for col in range(2):
            cell = outer.cell(0, col)
            nested = cell.add_table(rows=2, cols=2)
            nested.cell(0, 0).text = "Parameter"
            nested.cell(0, 1).text = "Value"
            nested.cell(1, 0).text = f"Field{col + 1}"
            nested.cell(1, 1).text = f"{col + 1}"
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "nested.docx"
            document.save(path)
            blocks, _items, cells = extract_docx(path)
        # 嵌套表作为父块 nested_tables 字段，其 items/cells 进扁平清单
        nested_table_ids = [
            entry["table_id"]
            for block in blocks
            for entry in (block.get("nested_tables") or [])
        ]
        self.assertEqual(
            sorted(nested_table_ids), ["TBL-000001-N001", "TBL-000001-N002"]
        )
        # 全部 cell_id 唯一——碰撞会在这里暴露
        all_cell_ids = [str(cell.get("cell_id") or "") for cell in cells]
        self.assertEqual(len(all_cell_ids), len(set(all_cell_ids)))


if __name__ == "__main__":
    unittest.main()
