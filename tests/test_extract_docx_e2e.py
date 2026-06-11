from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from atomize import build_atomic_candidates, build_chunks, extract_docx, mark_doc_regions


class ExtractDocxE2ETests(unittest.TestCase):
    def test_extract_docx_preserves_sections_and_generates_key_candidates(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "synthetic_standard.docx"
            self.write_synthetic_docx(input_path)

            blocks, table_items = extract_docx(input_path)
            mark_doc_regions(blocks, table_items)
            chunks = build_chunks(blocks, target_chars=800, include_regions={"body"})
            candidates = build_atomic_candidates(blocks, table_items, include_regions={"body"})

        blocks_by_text = {block["text"]: block for block in blocks}
        list_block = blocks_by_text["Display all segments"]

        self.assertEqual(blocks_by_text["Scope"]["type"], "heading")
        self.assertEqual(blocks_by_text["5.1 Security requirements"]["type"], "heading")
        self.assertEqual(list_block["type"], "paragraph")
        self.assertEqual(list_block["section_path"], ["Scope", "5.1 Security requirements"])
        self.assertTrue(all(block.get("doc_region") == "body" for block in blocks if block["text"] != "ABNT 2022 all rights reserved"))
        self.assertTrue(chunks)

        candidate_types = {candidate["requirement_type"] for candidate in candidates}
        self.assertIn("capability_matrix", candidate_types)
        self.assertIn("cosem_attribute_access", candidate_types)
        self.assertIn("event_definition", candidate_types)

        matrix_item = next(item for item in table_items if item.get("matrix_facts"))
        self.assertTrue(matrix_item["matrix_facts"])

        cosem_candidate = next(candidate for candidate in candidates if candidate["requirement_type"] == "cosem_attribute_access")
        self.assertIn("Clock.logical_name", cosem_candidate["object"])
        self.assertEqual(cosem_candidate["parameters"]["cosem_object"]["object_name"], "Clock")

        event_candidate = next(candidate for candidate in candidates if candidate["requirement_type"] == "event_definition")
        self.assertEqual(event_candidate["object"], "Event G1-SG10-E1")

    def write_synthetic_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("ABNT 2022 all rights reserved")
        document.add_heading("Scope", level=1)
        document.add_paragraph("The meter shall support xDLMS GET service.")
        document.add_paragraph("5.1 Security requirements")
        list_item = document.add_paragraph("Display all segments", style="List Paragraph")
        add_numbering(list_item)

        document.add_paragraph("Table 1 - Services xDLMS")
        matrix = document.add_table(rows=3, cols=3)
        matrix.cell(0, 0).text = "Customer application process"
        matrix.cell(0, 1).text = "xDLMS Service"
        matrix.cell(0, 2).text = "xDLMS Service"
        matrix.cell(1, 0).text = "Customer application process"
        matrix.cell(1, 1).text = "GET"
        matrix.cell(1, 2).text = "ACTION"
        matrix.cell(2, 0).text = "Public customer"
        matrix.cell(2, 1).text = "X"
        matrix.cell(2, 2).text = ""

        document.add_paragraph("Table 2 - COSEM Clock object")
        cosem = document.add_table(rows=3, cols=6)
        headers = ["#", "Object/attribute name", "CL", "Type", "Value", "Access rights RC/PC/SC/LC"]
        for index, header in enumerate(headers):
            cosem.cell(0, index).text = header
        for index, value in enumerate(["1", "Clock", "8", "", "0-0:1.0.0.255", ""]):
            cosem.cell(1, index).text = value
        for index, value in enumerate(["2", "logical_name", "", "octet-string[6]", "0000010000FF", "R-/--/R-/RW"]):
            cosem.cell(2, index).text = value

        document.add_paragraph("Table 3 - Events")
        events = document.add_table(rows=2, cols=4)
        event_headers = ["Group number", "Subgroup number", "Event number", "Description of the event"]
        for index, header in enumerate(event_headers):
            events.cell(0, index).text = header
        for index, value in enumerate(["1", "10", "1", "Power down"]):
            events.cell(1, index).text = value

        document.save(path)


def add_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


if __name__ == "__main__":
    unittest.main()
