"""合成需求文档测试语料（无需真实文档）——把真实文档实证过的关键场景固化为仓库用例。

场景出处（全部实证于 2026-07 真实文档评测）：
- STO/俄标：大参数表逐行、合并单元格展开、分组标题行、节号单元格、术语表边界
- SBD：技术规格矩阵（Clause|Requirement|Compliant）与商务表（价格/Lot）的表型区分
- EN 16314：清单逐项成块
- 审核 P1：映射/矩阵表 Phase 3 后置（判 other 不炸开）
"""
from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from ai_extract import (
    _supplement_parameter_table_rows,
    classify_table_kind,
)
from atomize import run_atomizer_pipeline
from io_utils import read_jsonl


def _build_corpus_docx(path: Path) -> None:
    document = Document()
    document.add_heading("5. Technical Requirements", level=1)

    # 场景 1：参数表（STO 型，含合并单元格展开、节号单元格、分组标题行）
    document.add_paragraph("Table 1 - Meter parameters")
    table = document.add_table(rows=6, cols=3)
    for i, h in enumerate(["No.", "Parameter Name", "Technical requirements"]):
        table.cell(0, i).text = h
    # 分组标题行（合并单元格 → 全同值）
    table.cell(1, 0).text = "1. GENERAL"
    table.cell(1, 1).text = "1. GENERAL"
    table.cell(1, 2).text = "1. GENERAL"
    # 节号 + 名称 + 要求 行
    table.cell(2, 0).text = "1.1"
    table.cell(2, 1).text = "Rated voltage"
    table.cell(2, 2).text = "The meter shall operate at 230 V nominal."
    table.cell(3, 0).text = "1.2"
    table.cell(3, 1).text = "Rated frequency"
    table.cell(3, 2).text = "The meter shall operate at 50 Hz nominal."
    # 稀疏行（单格有内容）
    table.cell(4, 0).text = "1.3"
    table.cell(4, 1).text = ""
    table.cell(4, 2).text = ""
    table.cell(5, 0).text = "1.4"
    table.cell(5, 1).text = "Backup power"
    table.cell(5, 2).text = "The meter shall include a reserve power supply."

    # 场景 2：术语表（不该出需求）
    document.add_paragraph("Table 2 - Terms and definitions")
    terms = document.add_table(rows=3, cols=3)
    for i, h in enumerate(["No.", "Term", "Definition"]):
        terms.cell(0, i).text = h
    terms.cell(1, 0).text = "1."
    terms.cell(1, 1).text = "Firmware"
    terms.cell(1, 2).text = "Software that processes information."
    terms.cell(2, 0).text = "2."
    terms.cell(2, 1).text = "Data"
    terms.cell(2, 2).text = "Information from measuring instruments."

    # 场景 3：商务表（含 Value 但不是需求表——价格/Lot 清单不得误判参数表）
    document.add_paragraph("Table 3 - Price schedule")
    price = document.add_table(rows=3, cols=3)
    for i, h in enumerate(["Lot#", "Description", "Quantity"]):
        price.cell(0, i).text = h
    price.cell(1, 0).text = "1"
    price.cell(1, 1).text = "Single phase meter"
    price.cell(1, 2).text = "500"
    price.cell(2, 0).text = "2"
    price.cell(2, 1).text = "Three phase meter"
    price.cell(2, 2).text = "250"

    # 场景 4：映射矩阵（行=OBIS 对象,列=访问级,格=权限事实——Phase 3 后置:判 other）
    document.add_paragraph("Table 4 - Access matrix")
    matrix = document.add_table(rows=3, cols=3)
    for i, h in enumerate(["Object", "Read", "Write"]):
        matrix.cell(0, i).text = h
    matrix.cell(1, 0).text = "0-0:1.0.0.255"
    matrix.cell(1, 1).text = "X"
    matrix.cell(1, 2).text = ""
    matrix.cell(2, 0).text = "0-0:96.1.0.255"
    matrix.cell(2, 1).text = "X"
    matrix.cell(2, 2).text = "X"

    # 场景 5：清单逐项成块（EN 型）
    document.add_paragraph("The meter shall satisfy the following:")
    for item in ("a) The meter shall store daily profiles.",
                 "b) The meter shall record power-down events.",
                 "c) The meter shall expose an optical port."):
        document.add_paragraph(item)

    document.save(path)


class SyntheticCorpusTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls._tmp = tempfile.TemporaryDirectory()
        cls.out = Path(cls._tmp.name)
        docx = cls.out / "corpus.docx"
        _build_corpus_docx(docx)
        run_atomizer_pipeline(docx, cls.out / "run", kb_paths=[])
        cls.blocks = read_jsonl(cls.out / "run" / "blocks.jsonl")
        cls.tables = [b for b in cls.blocks if str(b.get("type") or "") == "table"]

    @classmethod
    def tearDownClass(cls) -> None:
        cls._tmp.cleanup()

    def _table_by_title(self, needle: str) -> dict:
        for b in self.tables:
            if needle in str(b.get("table_title") or ""):
                return b
        raise AssertionError(f"table not found: {needle}")

    def test_parameter_table_qualified_and_expanded_per_row(self) -> None:
        """场景 1：参数表判 parameter；行级展开只出实质数据行（分组标题/稀疏行跳过）。"""
        block = self._table_by_title("Table 1")
        self.assertEqual(classify_table_kind(block), "parameter")

        rows = _supplement_parameter_table_rows([], [block])

        ids = [r["ai_req_id"] for r in rows]
        self.assertEqual(len(rows), 3)   # 1.1/1.2/1.4（1. GENERAL 与 1.3 稀疏行不进）
        self.assertEqual(
            ids,
            [f"PROW-DET-{block['block_id']}-R0002",
             f"PROW-DET-{block['block_id']}-R0003",
             f"PROW-DET-{block['block_id']}-R0005"],
        )
        first = rows[0]
        self.assertEqual(first["title"], "Rated voltage")
        self.assertEqual(first["source_quote"], "1.1 | Rated voltage | The meter shall operate at 230 V nominal.")
        self.assertEqual(first["source_mapping"], "deterministic_fallback")

    def test_expanded_quotes_are_verbatim_in_block_text(self) -> None:
        block = self._table_by_title("Table 1")
        rows = _supplement_parameter_table_rows([], [block])
        for r in rows:
            self.assertIn(r["source_quote"], str(block.get("text") or ""))

    def test_terms_table_not_qualified(self) -> None:
        """场景 2：术语表判 other（裁定：术语行不是需求）。"""
        block = self._table_by_title("Table 2")
        self.assertEqual(classify_table_kind(block), "other")
        self.assertEqual(_supplement_parameter_table_rows([], [block]), [])

    def test_price_schedule_not_misclassified(self) -> None:
        """场景 3：商务表（Lot/价格）不误判参数表——表头扩展正则的误伤护栏。"""
        block = self._table_by_title("Table 3")
        self.assertEqual(classify_table_kind(block), "other")

    def test_mapping_matrix_is_cell_mode_not_expanded(self) -> None:
        """场景 4：映射矩阵（table-structure-v2）→ mapping_matrix + cell 闭环，
        不走参数行展开（行只作容器，marker 格各自成 claim）。"""
        block = self._table_by_title("Table 4")
        self.assertEqual(classify_table_kind(block), "mapping_matrix")
        self.assertEqual(_supplement_parameter_table_rows([], [block]), [])
        self.assertEqual(str(block.get("leaf_mode") or ""), "cell")

    def test_list_items_become_separate_blocks(self) -> None:
        """场景 5：清单 a)/b)/c) 逐项独立成块（EN 型分块）。"""
        item_texts = [
            str(b.get("text") or "")
            for b in self.blocks
            if str(b.get("text") or "").strip().startswith(("a)", "b)", "c)"))
        ]
        self.assertEqual(len(item_texts), 3)
        self.assertTrue(any("daily profiles" in t for t in item_texts))
        self.assertTrue(any("optical port" in t for t in item_texts))

    def test_covered_rows_not_duplicated(self) -> None:
        """LLM 已覆盖的行不重复补（key_cell 命中判定）。"""
        block = self._table_by_title("Table 1")
        llm_req = {
            "ai_req_id": "AIR-1",
            "title": "电压规格",
            "description": "The meter shall operate at 230 V nominal.",
            "source_quote": "The meter shall operate at 230 V nominal.",
            "source_block_ids": [block["block_id"]],
        }
        rows = _supplement_parameter_table_rows([llm_req], [block])
        self.assertEqual(len(rows), 3)   # 1 LLM + 2 补行（1.2/1.4）,1.1 不重复
        ids = [r["ai_req_id"] for r in rows]
        self.assertEqual(ids.count("PROW-DET-%s-R0002" % block["block_id"]), 0)


if __name__ == "__main__":
    unittest.main()
